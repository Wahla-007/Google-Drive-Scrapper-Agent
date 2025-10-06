# # # # from __future__ import print_function
# # # # import os
# # # # import io
# # # # import psycopg2
# # # # from urllib.parse import quote_plus
# # # # from google.oauth2.credentials import Credentials
# # # # from google_auth_oauthlib.flow import InstalledAppFlow
# # # # from google.auth.transport.requests import Request
# # # # from googleapiclient.discovery import build
# # # # from datetime import datetime
# # # # import uuid

# # # # # Embedding + Vector DB
# # # # from openai import OpenAI
# # # # from qdrant_client import QdrantClient
# # # # from qdrant_client.http import models as qmodels

# # # # # PDF parsing
# # # # from PyPDF2 import PdfReader

# # # # # Configuration
# # # # import config

# # # # # Google Drive scopes (read & write)
# # # # SCOPES = [
# # # #     'https://www.googleapis.com/auth/drive.readonly',  # read all files
# # # #     'https://www.googleapis.com/auth/drive.file'       # read/write files created or opened by app
# # # # ]

# # # # # Initialize OpenAI and Qdrant
# # # # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # # # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)  # Increased timeout
# # # # COLLECTION_NAME = "fileData"  # your Qdrant collection name

# # # # # Batch size for Qdrant upsert
# # # # BATCH_SIZE = 100  # Process 100 points at a time


# # # # def authenticate():
# # # #     creds = None
# # # #     if os.path.exists(config.TOKEN_FILE):
# # # #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# # # #         print("Loaded existing token.json")

# # # #     if not creds or not creds.valid:
# # # #         if creds and creds.expired and creds.refresh_token:
# # # #             creds.refresh(Request())
# # # #             print("Token refreshed")
# # # #         else:
# # # #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# # # #             creds = flow.run_local_server(port=0)

# # # #         with open(config.TOKEN_FILE, 'w') as token_file:
# # # #             token_file.write(creds.to_json())
# # # #         print("token.json saved successfully")

# # # #     return build('drive', 'v3', credentials=creds)


# # # # def ensure_qdrant_collection():
# # # #     if not qdrant.collection_exists(COLLECTION_NAME):
# # # #         qdrant.create_collection(
# # # #             collection_name=COLLECTION_NAME,
# # # #             vectors_config=qmodels.VectorParams(
# # # #                 size=1536,
# # # #                 distance="Cosine"
# # # #             )
# # # #         )
# # # #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # # # def generate_embedding(text: str):
# # # #     response = openai_client.embeddings.create(
# # # #         model="text-embedding-3-small",
# # # #         input=text
# # # #     )
# # # #     return response.data[0].embedding


# # # # def drive_id_to_uuid(drive_id: str) -> str:
# # # #     """Convert Google Drive ID to a valid UUID for Qdrant."""
# # # #     # Use UUID v5 with a namespace to create deterministic UUIDs
# # # #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # # def extract_file_content(service, file_id, mime_type, file_name=""):
# # # #     """Download and extract text content depending on file type."""
# # # #     try:
# # # #         if mime_type == "application/vnd.google-apps.document":
# # # #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# # # #             return doc.decode("utf-8")

# # # #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# # # #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# # # #             return sheet.decode("utf-8")

# # # #         elif mime_type == "application/pdf":
# # # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # # #             reader = PdfReader(fh)
# # # #             return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

# # # #         elif mime_type.startswith("text/"):
# # # #             text = service.files().get_media(fileId=file_id).execute()
# # # #             content = text.decode("utf-8")
# # # #             # Check if file is empty
# # # #             if not content or content.strip() == "":
# # # #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# # # #                 return None
# # # #             return content

# # # #         else:
# # # #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# # # #             return None
# # # #     except Exception as e:
# # # #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# # # #         return None


# # # # def sync_drive(service, conn, root_folder_id):
# # # #     current_nodes = {}

# # # #     def walk_folder(folder_id, parent_id=None, path=""):
# # # #         query = f"'{folder_id}' in parents and trashed=false"
# # # #         results = service.files().list(
# # # #             q=query,
# # # #             spaces='drive',
# # # #             includeItemsFromAllDrives=True,
# # # #             supportsAllDrives=True,
# # # #             fields="files(id, name, mimeType, modifiedTime)"
# # # #         ).execute()
# # # #         items = results.get('files', [])

# # # #         for item in items:
# # # #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# # # #             mime_type = None if node_type == "folder" else item['mimeType']
# # # #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# # # #             node_path = f"{path}.{safe_name}" if path else safe_name

# # # #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# # # #             current_nodes[item['id']] = {
# # # #                 'name': item['name'],
# # # #                 'parent_id': parent_id,
# # # #                 'type': node_type,
# # # #                 'mime_type': mime_type,
# # # #                 'path': node_path,
# # # #                 'modified_time': modified_time
# # # #             }

# # # #             if node_type == "folder":
# # # #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# # # #     walk_folder(root_folder_id)

# # # #     # --- Postgres Sync ---
# # # #     with conn.cursor() as cur:
# # # #         for gid, info in current_nodes.items():
# # # #             cur.execute("""
# # # #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# # # #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# # # #                 ON CONFLICT (id) DO UPDATE
# # # #                 SET name = EXCLUDED.name,
# # # #                     parent_id = EXCLUDED.parent_id,
# # # #                     type = EXCLUDED.type,
# # # #                     mime_type = EXCLUDED.mime_type,
# # # #                     path = EXCLUDED.path,
# # # #                     modified_time = EXCLUDED.modified_time
# # # #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# # # #             """, (gid, info['name'], info['parent_id'], info['type'],
# # # #                   info['mime_type'], info['path'], info['modified_time']))

# # # #         if current_nodes:
# # # #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# # # #         else:
# # # #             cur.execute("DELETE FROM nodes;")
# # # #         conn.commit()

# # # #     # --- Qdrant Sync ---
# # # #     ensure_qdrant_collection()
# # # #     points = []

# # # #     for gid, info in current_nodes.items():
# # # #         text_content = None
# # # #         if info['type'] == "file" and info['mime_type']:
# # # #             text_content = extract_file_content(service, gid, info['mime_type'], info['name'])

# # # #         if text_content is None:
# # # #             # Skip files we couldn't read (empty, unsupported, or errored)
# # # #             continue

# # # #         # --- Convert Drive ID to deterministic UUID for Qdrant ---
# # # #         point_id = drive_id_to_uuid(gid)
# # # #         content_to_embed = text_content[:8000]  # limit to 8k chars
        
# # # #         try:
# # # #             embedding = generate_embedding(content_to_embed)
# # # #         except Exception as e:
# # # #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# # # #             continue

# # # #         points.append(qmodels.PointStruct(
# # # #             id=point_id,
# # # #             vector=embedding,
# # # #             payload={
# # # #                 "drive_id": gid,  # original Google Drive ID
# # # #                 "name": info['name'],
# # # #                 "type": info['type'],
# # # #                 "mime_type": info['mime_type'],
# # # #                 "path": info['path'],
# # # #                 "modified_time": info['modified_time'].isoformat(),
# # # #                 "content": text_content[:10000]  # Store limited content in payload
# # # #             }
# # # #         ))

# # # #     # Upload to Qdrant in batches to avoid timeout
# # # #     if points:
# # # #         total_points = len(points)
# # # #         print(f"📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# # # #         for i in range(0, total_points, BATCH_SIZE):
# # # #             batch = points[i:i + BATCH_SIZE]
# # # #             try:
# # # #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# # # #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# # # #             except Exception as e:
# # # #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# # # #                 # Continue with next batch instead of failing completely
# # # #                 continue
        
# # # #         print(f"✅ Successfully synced {total_points} items to Qdrant")
# # # #     else:
# # # #         print("⚠️ No points to upload to Qdrant")


# # # # if __name__ == '__main__':
# # # #     service = authenticate()
# # # #     conn = psycopg2.connect(config.DATABASE_URL)

# # # #     print("Syncing Google Drive → Postgres + Qdrant (with file contents)...")
# # # #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# # # #     print("Done! DB + Qdrant are now consistent with Drive.")

# # # #     conn.close()

# # # from __future__ import print_function
# # # import os
# # # import io
# # # import psycopg2
# # # from urllib.parse import quote_plus
# # # from google.oauth2.credentials import Credentials
# # # from google_auth_oauthlib.flow import InstalledAppFlow
# # # from google.auth.transport.requests import Request
# # # from googleapiclient.discovery import build
# # # from datetime import datetime
# # # import uuid

# # # # Embedding + Vector DB
# # # from openai import OpenAI
# # # from qdrant_client import QdrantClient
# # # from qdrant_client.http import models as qmodels

# # # # PDF parsing
# # # from PyPDF2 import PdfReader

# # # # Document parsing
# # # from docx import Document  # python-docx for Word documents
# # # import pytesseract  # OCR for images
# # # from PIL import Image  # Image processing

# # # # Configuration
# # # import config

# # # # Google Drive scopes (read & write)
# # # SCOPES = [
# # #     'https://www.googleapis.com/auth/drive.readonly',  # read all files
# # #     'https://www.googleapis.com/auth/drive.file'       # read/write files created or opened by app
# # # ]

# # # # Initialize OpenAI and Qdrant
# # # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)  # Increased timeout
# # # COLLECTION_NAME = "fileData"  # your Qdrant collection name

# # # # Batch size for Qdrant upsert
# # # BATCH_SIZE = 100  # Process 100 points at a time


# # # def authenticate():
# # #     creds = None
# # #     if os.path.exists(config.TOKEN_FILE):
# # #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# # #         print("Loaded existing token.json")

# # #     if not creds or not creds.valid:
# # #         if creds and creds.expired and creds.refresh_token:
# # #             creds.refresh(Request())
# # #             print("Token refreshed")
# # #         else:
# # #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# # #             creds = flow.run_local_server(port=0)

# # #         with open(config.TOKEN_FILE, 'w') as token_file:
# # #             token_file.write(creds.to_json())
# # #         print("token.json saved successfully")

# # #     return build('drive', 'v3', credentials=creds)


# # # def ensure_qdrant_collection():
# # #     if not qdrant.collection_exists(COLLECTION_NAME):
# # #         qdrant.create_collection(
# # #             collection_name=COLLECTION_NAME,
# # #             vectors_config=qmodels.VectorParams(
# # #                 size=1536,
# # #                 distance="Cosine"
# # #             )
# # #         )
# # #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # # def generate_embedding(text: str):
# # #     response = openai_client.embeddings.create(
# # #         model="text-embedding-3-small",
# # #         input=text
# # #     )
# # #     return response.data[0].embedding


# # # def drive_id_to_uuid(drive_id: str) -> str:
# # #     """Convert Google Drive ID to a valid UUID for Qdrant."""
# # #     # Use UUID v5 with a namespace to create deterministic UUIDs
# # #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # def extract_file_content(service, file_id, mime_type, file_name=""):
# # #     """Download and extract text content depending on file type."""
# # #     try:
# # #         # Google Docs
# # #         if mime_type == "application/vnd.google-apps.document":
# # #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# # #             return doc.decode("utf-8")

# # #         # Google Sheets
# # #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# # #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# # #             return sheet.decode("utf-8")

# # #         # PDF files
# # #         elif mime_type == "application/pdf":
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             reader = PdfReader(fh)
# # #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ PDF '{file_name}' has no extractable text")
# # #                 return None
# # #             return text

# # #         # Word Documents (.docx)
# # #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             doc = Document(fh)
# # #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ Word document '{file_name}' is empty")
# # #                 return None
# # #             return text

# # #         # Images (JPEG, PNG, GIF, BMP, TIFF) - use OCR
# # #         elif mime_type.startswith("image/"):
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             image = Image.open(fh)
            
# # #             # Perform OCR
# # #             text = pytesseract.image_to_string(image)
            
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ Image '{file_name}' contains no readable text (OCR found nothing)")
# # #                 return None
            
# # #             return text.strip()

# # #         # Plain text files
# # #         elif mime_type.startswith("text/"):
# # #             text = service.files().get_media(fileId=file_id).execute()
# # #             content = text.decode("utf-8")
# # #             # Check if file is empty
# # #             if not content or content.strip() == "":
# # #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# # #                 return None
# # #             return content

# # #         else:
# # #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# # #             return None
            
# # #     except Exception as e:
# # #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# # #         return None


# # # def sync_drive(service, conn, root_folder_id):
# # #     current_nodes = {}
# # #     stats = {
# # #         'new_files': [],
# # #         'updated_files': [],
# # #         'unchanged_files': [],
# # #         'deleted_files': [],
# # #         'skipped_files': []
# # #     }

# # #     def walk_folder(folder_id, parent_id=None, path=""):
# # #         query = f"'{folder_id}' in parents and trashed=false"
# # #         results = service.files().list(
# # #             q=query,
# # #             spaces='drive',
# # #             includeItemsFromAllDrives=True,
# # #             supportsAllDrives=True,
# # #             fields="files(id, name, mimeType, modifiedTime)"
# # #         ).execute()
# # #         items = results.get('files', [])

# # #         for item in items:
# # #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# # #             mime_type = None if node_type == "folder" else item['mimeType']
# # #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# # #             node_path = f"{path}.{safe_name}" if path else safe_name

# # #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# # #             current_nodes[item['id']] = {
# # #                 'name': item['name'],
# # #                 'parent_id': parent_id,
# # #                 'type': node_type,
# # #                 'mime_type': mime_type,
# # #                 'path': node_path,
# # #                 'modified_time': modified_time
# # #             }

# # #             if node_type == "folder":
# # #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# # #     walk_folder(root_folder_id)

# # #     # --- Postgres Sync ---
# # #     with conn.cursor() as cur:
# # #         # Get existing files to track changes
# # #         cur.execute("SELECT id, modified_time FROM nodes WHERE type = 'file'")
# # #         existing_files = {row[0]: row[1] for row in cur.fetchall()}
        
# # #         for gid, info in current_nodes.items():
# # #             # Check if file is new or updated
# # #             if gid not in existing_files:
# # #                 if info['type'] == 'file':
# # #                     stats['new_files'].append(info['name'])
# # #             elif info['type'] == 'file' and existing_files[gid] != info['modified_time']:
# # #                 stats['updated_files'].append(info['name'])
# # #             elif info['type'] == 'file':
# # #                 stats['unchanged_files'].append(info['name'])
            
# # #             cur.execute("""
# # #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# # #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# # #                 ON CONFLICT (id) DO UPDATE
# # #                 SET name = EXCLUDED.name,
# # #                     parent_id = EXCLUDED.parent_id,
# # #                     type = EXCLUDED.type,
# # #                     mime_type = EXCLUDED.mime_type,
# # #                     path = EXCLUDED.path,
# # #                     modified_time = EXCLUDED.modified_time
# # #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# # #             """, (gid, info['name'], info['parent_id'], info['type'],
# # #                   info['mime_type'], info['path'], info['modified_time']))

# # #         # Track deleted files
# # #         if current_nodes:
# # #             current_file_ids = [gid for gid, info in current_nodes.items() if info['type'] == 'file']
# # #             deleted_ids = set(existing_files.keys()) - set(current_file_ids)
# # #             if deleted_ids:
# # #                 cur.execute("SELECT name FROM nodes WHERE id IN %s", (tuple(deleted_ids),))
# # #                 stats['deleted_files'] = [row[0] for row in cur.fetchall()]
            
# # #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# # #         else:
# # #             cur.execute("SELECT name FROM nodes WHERE type = 'file'")
# # #             stats['deleted_files'] = [row[0] for row in cur.fetchall()]
# # #             cur.execute("DELETE FROM nodes;")
# # #         conn.commit()

# # #     # --- Qdrant Sync ---
# # #     ensure_qdrant_collection()
    
# # #     # Get existing points from Qdrant to track what's new/updated
# # #     try:
# # #         existing_points = qdrant.scroll(
# # #             collection_name=COLLECTION_NAME,
# # #             limit=10000,
# # #             with_payload=True,
# # #             with_vectors=False
# # #         )[0]
# # #         existing_drive_ids = {point.payload.get('drive_id'): point.payload.get('modified_time') 
# # #                               for point in existing_points}
# # #     except Exception as e:
# # #         print(f"⚠️ Could not fetch existing Qdrant points: {e}")
# # #         existing_drive_ids = {}
    
# # #     points = []
# # #     qdrant_stats = {
# # #         'new': [],
# # #         'updated': [],
# # #         'skipped': []
# # #     }

# # #     for gid, info in current_nodes.items():
# # #         text_content = None
# # #         if info['type'] == "file" and info['mime_type']:
# # #             text_content = extract_file_content(service, gid, info['mime_type'], info['name'])

# # #         if text_content is None:
# # #             # Skip files we couldn't read (empty, unsupported, or errored)
# # #             stats['skipped_files'].append(info['name'])
# # #             continue
        
# # #         # Check if this is new or updated in Qdrant
# # #         is_new = gid not in existing_drive_ids
# # #         is_updated = (not is_new and 
# # #                      existing_drive_ids[gid] != info['modified_time'].isoformat())
        
# # #         if is_new:
# # #             qdrant_stats['new'].append(info['name'])
# # #         elif is_updated:
# # #             qdrant_stats['updated'].append(info['name'])
# # #         else:
# # #             # File exists and hasn't changed, but we'll re-upload to ensure consistency
# # #             pass

# # #         # --- Convert Drive ID to deterministic UUID for Qdrant ---
# # #         point_id = drive_id_to_uuid(gid)
# # #         content_to_embed = text_content[:8000]  # limit to 8k chars
        
# # #         try:
# # #             embedding = generate_embedding(content_to_embed)
# # #         except Exception as e:
# # #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# # #             continue

# # #         points.append(qmodels.PointStruct(
# # #             id=point_id,
# # #             vector=embedding,
# # #             payload={
# # #                 "drive_id": gid,  # original Google Drive ID
# # #                 "name": info['name'],
# # #                 "type": info['type'],
# # #                 "mime_type": info['mime_type'],
# # #                 "path": info['path'],
# # #                 "modified_time": info['modified_time'].isoformat(),
# # #                 "content": text_content[:10000]  # Store limited content in payload
# # #             }
# # #         ))

# # #     # Upload to Qdrant in batches to avoid timeout
# # #     if points:
# # #         total_points = len(points)
# # #         print(f"\n📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# # #         for i in range(0, total_points, BATCH_SIZE):
# # #             batch = points[i:i + BATCH_SIZE]
# # #             try:
# # #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# # #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# # #             except Exception as e:
# # #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# # #                 # Continue with next batch instead of failing completely
# # #                 continue
        
# # #         print(f"✅ Successfully synced {total_points} items to Qdrant")
# # #     else:
# # #         print("⚠️ No points to upload to Qdrant")
    
# # #     # Print detailed statistics
# # #     print("\n" + "="*60)
# # #     print("📊 SYNC SUMMARY")
# # #     print("="*60)
    
# # #     print("\n🗄️  PostgreSQL Changes:")
# # #     if stats['new_files']:
# # #         print(f"  ✨ New files ({len(stats['new_files'])}):")
# # #         for name in stats['new_files'][:10]:  # Show first 10
# # #             print(f"     • {name}")
# # #         if len(stats['new_files']) > 10:
# # #             print(f"     ... and {len(stats['new_files']) - 10} more")
    
# # #     if stats['updated_files']:
# # #         print(f"  🔄 Updated files ({len(stats['updated_files'])}):")
# # #         for name in stats['updated_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['updated_files']) > 10:
# # #             print(f"     ... and {len(stats['updated_files']) - 10} more")
    
# # #     if stats['deleted_files']:
# # #         print(f"  🗑️  Deleted files ({len(stats['deleted_files'])}):")
# # #         for name in stats['deleted_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['deleted_files']) > 10:
# # #             print(f"     ... and {len(stats['deleted_files']) - 10} more")
    
# # #     if stats['unchanged_files']:
# # #         print(f"  ✓ Unchanged files: {len(stats['unchanged_files'])}")
    
# # #     print("\n🔍 Qdrant Vector Database Changes:")
# # #     if qdrant_stats['new']:
# # #         print(f"  ✨ New embeddings ({len(qdrant_stats['new'])}):")
# # #         for name in qdrant_stats['new'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['new']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['new']) - 10} more")
    
# # #     if qdrant_stats['updated']:
# # #         print(f"  🔄 Updated embeddings ({len(qdrant_stats['updated'])}):")
# # #         for name in qdrant_stats['updated'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['updated']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['updated']) - 10} more")
    
# # #     if stats['skipped_files']:
# # #         print(f"\n⚠️  Skipped files ({len(stats['skipped_files'])}):")
# # #         for name in stats['skipped_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['skipped_files']) > 10:
# # #             print(f"     ... and {len(stats['skipped_files']) - 10} more")
    
# # #     print("\n" + "="*60)


# # # if __name__ == '__main__':
# # #     service = authenticate()
# # #     conn = psycopg2.connect(config.DATABASE_URL)

# # #     print("Syncing Google Drive → Postgres + Qdrant (with file contents)...")
# # #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# # #     print("Done! DB + Qdrant are now consistent with Drive.")

# # #     conn.close()
# # from __future__ import print_function
# # import os
# # import io
# # import psycopg2
# # from urllib.parse import quote_plus
# # from google.oauth2.credentials import Credentials
# # from google_auth_oauthlib.flow import InstalledAppFlow
# # from google.auth.transport.requests import Request
# # from googleapiclient.discovery import build
# # from datetime import datetime
# # import uuid

# # # Embedding + Vector DB
# # from openai import OpenAI
# # from qdrant_client import QdrantClient
# # from qdrant_client.http import models as qmodels

# # # PDF parsing
# # from PyPDF2 import PdfReader

# # # Document parsing
# # from docx import Document  # python-docx for Word documents
# # import pytesseract  # OCR for images
# # from PIL import Image  # Image processing

# # # Configuration
# # import config

# # # Google Drive scopes (read & write)
# # SCOPES = [
# #     'https://www.googleapis.com/auth/drive.readonly',  # read all files
# #     'https://www.googleapis.com/auth/drive.file'       # read/write files created or opened by app
# # ]

# # # Initialize OpenAI and Qdrant
# # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)
# # COLLECTION_NAME = "fileData"

# # # Batch size for Qdrant upsert
# # BATCH_SIZE = 100


# # def authenticate():
# #     creds = None
# #     if os.path.exists(config.TOKEN_FILE):
# #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# #         print("Loaded existing token.json")

# #     if not creds or not creds.valid:
# #         if creds and creds.expired and creds.refresh_token:
# #             creds.refresh(Request())
# #             print("Token refreshed")
# #         else:
# #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# #             creds = flow.run_local_server(port=0)

# #         with open(config.TOKEN_FILE, 'w') as token_file:
# #             token_file.write(creds.to_json())
# #         print("token.json saved successfully")

# #     return build('drive', 'v3', credentials=creds)


# # def ensure_qdrant_collection():
# #     if not qdrant.collection_exists(COLLECTION_NAME):
# #         qdrant.create_collection(
# #             collection_name=COLLECTION_NAME,
# #             vectors_config=qmodels.VectorParams(
# #                 size=1536,
# #                 distance="Cosine"
# #             )
# #         )
# #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # def generate_embedding(text: str):
# #     response = openai_client.embeddings.create(
# #         model="text-embedding-3-small",
# #         input=text
# #     )
# #     return response.data[0].embedding


# # def drive_id_to_uuid(drive_id: str) -> str:
# #     """Convert Google Drive ID to a valid UUID for Qdrant."""
# #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # def extract_file_content(service, file_id, mime_type, file_name=""):
# #     """Download and extract text content depending on file type."""
# #     try:
# #         # Google Docs
# #         if mime_type == "application/vnd.google-apps.document":
# #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# #             return doc.decode("utf-8")

# #         # Google Sheets
# #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# #             return sheet.decode("utf-8")

# #         # PDF files
# #         elif mime_type == "application/pdf":
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             reader = PdfReader(fh)
# #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# #             if not text or text.strip() == "":
# #                 print(f"⚠️ PDF '{file_name}' has no extractable text")
# #                 return None
# #             return text

# #         # Word Documents (.docx)
# #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             doc = Document(fh)
# #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# #             if not text or text.strip() == "":
# #                 print(f"⚠️ Word document '{file_name}' is empty")
# #                 return None
# #             return text

# #         # Images (JPEG, PNG, GIF, BMP, TIFF) - use OCR
# #         elif mime_type.startswith("image/"):
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             image = Image.open(fh)
            
# #             # Perform OCR
# #             text = pytesseract.image_to_string(image)
            
# #             if not text or text.strip() == "":
# #                 print(f"⚠️ Image '{file_name}' contains no readable text (OCR found nothing)")
# #                 return None
            
# #             return text.strip()

# #         # Plain text files
# #         elif mime_type.startswith("text/"):
# #             text = service.files().get_media(fileId=file_id).execute()
# #             content = text.decode("utf-8")
# #             # Check if file is empty
# #             if not content or content.strip() == "":
# #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# #                 return None
# #             return content

# #         else:
# #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# #             return None
            
# #     except Exception as e:
# #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# #         return None


# # def sync_drive(service, conn, root_folder_id):
# #     current_nodes = {}
# #     stats = {
# #         'new_files': [],
# #         'updated_files': [],
# #         'unchanged_files': [],
# #         'deleted_files': [],
# #         'skipped_files': []
# #     }

# #     def walk_folder(folder_id, parent_id=None, path=""):
# #         query = f"'{folder_id}' in parents and trashed=false"
# #         results = service.files().list(
# #             q=query,
# #             spaces='drive',
# #             includeItemsFromAllDrives=True,
# #             supportsAllDrives=True,
# #             fields="files(id, name, mimeType, modifiedTime)"
# #         ).execute()
# #         items = results.get('files', [])

# #         for item in items:
# #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# #             mime_type = None if node_type == "folder" else item['mimeType']
# #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# #             node_path = f"{path}.{safe_name}" if path else safe_name

# #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# #             current_nodes[item['id']] = {
# #                 'name': item['name'],
# #                 'parent_id': parent_id,
# #                 'type': node_type,
# #                 'mime_type': mime_type,
# #                 'path': node_path,
# #                 'modified_time': modified_time
# #             }

# #             if node_type == "folder":
# #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# #     walk_folder(root_folder_id)

# #     # --- Postgres Sync ---
# #     with conn.cursor() as cur:
# #         # Get existing files to track changes
# #         cur.execute("SELECT id, modified_time FROM nodes WHERE type = 'file'")
# #         existing_files = {row[0]: row[1] for row in cur.fetchall()}
        
# #         for gid, info in current_nodes.items():
# #             # Check if file is new or updated
# #             if gid not in existing_files:
# #                 if info['type'] == 'file':
# #                     stats['new_files'].append(info['name'])
# #             elif info['type'] == 'file' and existing_files[gid] != info['modified_time']:
# #                 stats['updated_files'].append(info['name'])
# #             elif info['type'] == 'file':
# #                 stats['unchanged_files'].append(info['name'])
            
# #             cur.execute("""
# #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# #                 ON CONFLICT (id) DO UPDATE
# #                 SET name = EXCLUDED.name,
# #                     parent_id = EXCLUDED.parent_id,
# #                     type = EXCLUDED.type,
# #                     mime_type = EXCLUDED.mime_type,
# #                     path = EXCLUDED.path,
# #                     modified_time = EXCLUDED.modified_time
# #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# #             """, (gid, info['name'], info['parent_id'], info['type'],
# #                   info['mime_type'], info['path'], info['modified_time']))

# #         # Track deleted files
# #         if current_nodes:
# #             current_file_ids = [gid for gid, info in current_nodes.items() if info['type'] == 'file']
# #             deleted_ids = set(existing_files.keys()) - set(current_file_ids)
# #             if deleted_ids:
# #                 cur.execute("SELECT name FROM nodes WHERE id IN %s", (tuple(deleted_ids),))
# #                 stats['deleted_files'] = [row[0] for row in cur.fetchall()]
            
# #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# #         else:
# #             cur.execute("SELECT name FROM nodes WHERE type = 'file'")
# #             stats['deleted_files'] = [row[0] for row in cur.fetchall()]
# #             cur.execute("DELETE FROM nodes;")
# #         conn.commit()

# #     # --- Qdrant Sync (INCREMENTAL) ---
# #     ensure_qdrant_collection()
    
# #     # Get existing points from Qdrant
# #     try:
# #         existing_points = qdrant.scroll(
# #             collection_name=COLLECTION_NAME,
# #             limit=10000,
# #             with_payload=True,
# #             with_vectors=False
# #         )[0]
# #         existing_drive_ids = {point.payload.get('drive_id'): {
# #             'modified_time': point.payload.get('modified_time'),
# #             'point_id': point.id
# #         } for point in existing_points}
# #     except Exception as e:
# #         print(f"⚠️ Could not fetch existing Qdrant points: {e}")
# #         existing_drive_ids = {}
    
# #     # Identify files that need action
# #     files_to_add = []  # New files
# #     files_to_update = []  # Modified files
# #     files_to_delete = []  # Deleted from Drive
    
# #     # Find files in Drive that need to be added or updated
# #     for gid, info in current_nodes.items():
# #         if info['type'] != "file" or not info['mime_type']:
# #             continue
            
# #         is_new = gid not in existing_drive_ids
# #         is_updated = (not is_new and 
# #                      existing_drive_ids[gid]['modified_time'] != info['modified_time'].isoformat())
        
# #         if is_new:
# #             files_to_add.append((gid, info))
# #         elif is_updated:
# #             files_to_update.append((gid, info))
    
# #     # Find files in Qdrant that have been deleted from Drive
# #     current_drive_file_ids = {gid for gid, info in current_nodes.items() if info['type'] == 'file'}
# #     for drive_id, point_info in existing_drive_ids.items():
# #         if drive_id not in current_drive_file_ids:
# #             files_to_delete.append((drive_id, point_info['point_id']))
    
# #     # Process new and updated files
# #     points_to_upsert = []
# #     qdrant_stats = {
# #         'new': [],
# #         'updated': [],
# #         'deleted': [],
# #         'skipped': []
# #     }
    
# #     print(f"\n🔄 Processing {len(files_to_add)} new files and {len(files_to_update)} updated files...")
    
# #     for gid, info in files_to_add + files_to_update:
# #         is_new_file = (gid, info) in files_to_add
        
# #         text_content = extract_file_content(service, gid, info['mime_type'], info['name'])
        
# #         if text_content is None:
# #             stats['skipped_files'].append(info['name'])
# #             continue
        
# #         if is_new_file:
# #             qdrant_stats['new'].append(info['name'])
# #         else:
# #             qdrant_stats['updated'].append(info['name'])
        
# #         point_id = drive_id_to_uuid(gid)
# #         content_to_embed = text_content[:8000]
        
# #         try:
# #             embedding = generate_embedding(content_to_embed)
# #         except Exception as e:
# #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# #             continue
        
# #         points_to_upsert.append(qmodels.PointStruct(
# #             id=point_id,
# #             vector=embedding,
# #             payload={
# #                 "drive_id": gid,
# #                 "name": info['name'],
# #                 "type": info['type'],
# #                 "mime_type": info['mime_type'],
# #                 "path": info['path'],
# #                 "modified_time": info['modified_time'].isoformat(),
# #                 "content": text_content[:10000]
# #             }
# #         ))
    
# #     # Upload new/updated points to Qdrant in batches
# #     if points_to_upsert:
# #         total_points = len(points_to_upsert)
# #         print(f"\n📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# #         for i in range(0, total_points, BATCH_SIZE):
# #             batch = points_to_upsert[i:i + BATCH_SIZE]
# #             try:
# #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# #             except Exception as e:
# #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# #                 continue
        
# #         print(f"✅ Successfully synced {total_points} items to Qdrant")
    
# #     # Delete removed files from Qdrant
# #     if files_to_delete:
# #         print(f"\n🗑️  Deleting {len(files_to_delete)} removed files from Qdrant...")
# #         point_ids_to_delete = [point_id for _, point_id in files_to_delete]
        
# #         try:
# #             qdrant.delete(
# #                 collection_name=COLLECTION_NAME,
# #                 points_selector=qmodels.PointIdsList(points=point_ids_to_delete)
# #             )
# #             qdrant_stats['deleted'] = stats['deleted_files']
# #             print(f"✅ Deleted {len(point_ids_to_delete)} points from Qdrant")
# #         except Exception as e:
# #             print(f"❌ Failed to delete points from Qdrant: {e}")
    
# #     # Print detailed statistics
# #     print("\n" + "="*60)
# #     print("📊 INCREMENTAL SYNC SUMMARY")
# #     print("="*60)
    
# #     print("\n🗄️  PostgreSQL Changes:")
# #     if stats['new_files']:
# #         print(f"  ✨ New files ({len(stats['new_files'])}):")
# #         for name in stats['new_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['new_files']) > 10:
# #             print(f"     ... and {len(stats['new_files']) - 10} more")
    
# #     if stats['updated_files']:
# #         print(f"  🔄 Updated files ({len(stats['updated_files'])}):")
# #         for name in stats['updated_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['updated_files']) > 10:
# #             print(f"     ... and {len(stats['updated_files']) - 10} more")
    
# #     if stats['deleted_files']:
# #         print(f"  🗑️  Deleted files ({len(stats['deleted_files'])}):")
# #         for name in stats['deleted_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['deleted_files']) > 10:
# #             print(f"     ... and {len(stats['deleted_files']) - 10} more")
    
# #     if stats['unchanged_files']:
# #         print(f"  ✓ Unchanged files: {len(stats['unchanged_files'])} (not processed)")
    
# #     print("\n🔍 Qdrant Vector Database Changes:")
# #     if qdrant_stats['new']:
# #         print(f"  ✨ New embeddings ({len(qdrant_stats['new'])}):")
# #         for name in qdrant_stats['new'][:10]:
# #             print(f"     • {name}")
# #         if len(qdrant_stats['new']) > 10:
# #             print(f"     ... and {len(qdrant_stats['new']) - 10} more")
    
# #     if qdrant_stats['updated']:
# #         print(f"  🔄 Updated embeddings ({len(qdrant_stats['updated'])}):")
# #         for name in qdrant_stats['updated'][:10]:
# #             print(f"     • {name}")
# #         if len(qdrant_stats['updated']) > 10:
# #             print(f"     ... and {len(qdrant_stats['updated']) - 10} more")
    
# #     if qdrant_stats['deleted']:
# #         print(f"  🗑️  Deleted embeddings ({len(qdrant_stats['deleted'])}):")
# #         for name in qdrant_stats['deleted'][:10]:
# #             print(f"     • {name}")
# #         if len(qdrant_stats['deleted']) > 10:
# #             print(f"     ... and {len(qdrant_stats['deleted']) - 10} more")
    
# #     if stats['skipped_files']:
# #         print(f"\n⚠️  Skipped files ({len(stats['skipped_files'])}):")
# #         for name in stats['skipped_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['skipped_files']) > 10:
# #             print(f"     ... and {len(stats['skipped_files']) - 10} more")
    
# #     if not (qdrant_stats['new'] or qdrant_stats['updated'] or qdrant_stats['deleted']):
# #         print(f"  ✓ No changes - all files are up to date!")
    
# #     print("\n" + "="*60)


# # if __name__ == '__main__':
# #     service = authenticate()
# #     conn = psycopg2.connect(config.DATABASE_URL)

# #     print("Syncing Google Drive → Postgres + Qdrant (INCREMENTAL MODE)...")
# #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# #     print("Done! Only changed files were processed.")

# #     conn.close()

# # from __future__ import print_function
# # import os
# # import io
# # import psycopg2
# # from urllib.parse import quote_plus
# # from google.oauth2.credentials import Credentials
# # from google_auth_oauthlib.flow import InstalledAppFlow
# # from google.auth.transport.requests import Request
# # from googleapiclient.discovery import build
# # from datetime import datetime
# # import uuid

# # # Embedding + Vector DB
# # from openai import OpenAI
# # from qdrant_client import QdrantClient
# # from qdrant_client.http import models as qmodels

# # # PDF parsing
# # from PyPDF2 import PdfReader

# # # Document parsing
# # from docx import Document
# # import pytesseract
# # from PIL import Image

# # # Configuration
# # import config

# # # Google Drive scopes
# # SCOPES = [
# #     'https://www.googleapis.com/auth/drive.readonly',
# #     'https://www.googleapis.com/auth/drive.file'
# # ]

# # # Initialize OpenAI and Qdrant
# # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)
# # COLLECTION_NAME = "fileData"

# # BATCH_SIZE = 100


# # def authenticate():
# #     creds = None
# #     if os.path.exists(config.TOKEN_FILE):
# #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# #         print("Loaded existing token.json")

# #     if not creds or not creds.valid:
# #         if creds and creds.expired and creds.refresh_token:
# #             creds.refresh(Request())
# #             print("Token refreshed")
# #         else:
# #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# #             creds = flow.run_local_server(port=0)

# #         with open(config.TOKEN_FILE, 'w') as token_file:
# #             token_file.write(creds.to_json())
# #         print("token.json saved successfully")

# #     return build('drive', 'v3', credentials=creds)


# # def ensure_qdrant_collection():
# #     if not qdrant.collection_exists(COLLECTION_NAME):
# #         qdrant.create_collection(
# #             collection_name=COLLECTION_NAME,
# #             vectors_config=qmodels.VectorParams(
# #                 size=1536,
# #                 distance="Cosine"
# #             )
# #         )
# #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # def generate_embedding(text: str):
# #     response = openai_client.embeddings.create(
# #         model="text-embedding-3-small",
# #         input=text
# #     )
# #     return response.data[0].embedding


# # def drive_id_to_uuid(drive_id: str) -> str:
# #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # def extract_file_content(service, file_id, mime_type, file_name=""):
# #     """Download and extract text content depending on file type."""
# #     try:
# #         # Google Docs
# #         if mime_type == "application/vnd.google-apps.document":
# #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# #             return doc.decode("utf-8")

# #         # Google Sheets
# #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# #             return sheet.decode("utf-8")

# #         # PDF files
# #         elif mime_type == "application/pdf":
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             reader = PdfReader(fh)
# #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# #             if not text or text.strip() == "":
# #                 print(f"⚠️ PDF '{file_name}' has no extractable text")
# #                 return None
# #             return text

# #         # Word Documents (.docx)
# #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             doc = Document(fh)
# #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# #             if not text or text.strip() == "":
# #                 print(f"⚠️ Word document '{file_name}' is empty")
# #                 return None
# #             return text

# #         # Images (JPEG, PNG, GIF, BMP, TIFF) - use OCR
# #         elif mime_type.startswith("image/"):
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             image = Image.open(fh)
            
# #             text = pytesseract.image_to_string(image)
            
# #             if not text or text.strip() == "":
# #                 print(f"⚠️ Image '{file_name}' contains no readable text (OCR found nothing)")
# #                 return None
            
# #             return text.strip()

# #         # Plain text files
# #         elif mime_type.startswith("text/"):
# #             text = service.files().get_media(fileId=file_id).execute()
# #             content = text.decode("utf-8")
# #             if not content or content.strip() == "":
# #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# #                 return None
# #             return content

# #         else:
# #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# #             return None
            
# #     except Exception as e:
# #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# #         return None


# # def sync_drive(service, conn, root_folder_id):
# #     current_nodes = {}
# #     stats = {
# #         'new_files': [],
# #         'updated_files': [],
# #         'unchanged_files': [],
# #         'deleted_files': [],
# #         'new_folders': [],
# #         'updated_folders': [],
# #         'deleted_folders': [],
# #         'skipped_files': []
# #     }

# #     def walk_folder(folder_id, parent_id=None, path=""):
# #         query = f"'{folder_id}' in parents and trashed=false"
# #         results = service.files().list(
# #             q=query,
# #             spaces='drive',
# #             includeItemsFromAllDrives=True,
# #             supportsAllDrives=True,
# #             fields="files(id, name, mimeType, modifiedTime)"
# #         ).execute()
# #         items = results.get('files', [])

# #         for item in items:
# #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# #             mime_type = None if node_type == "folder" else item['mimeType']
# #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# #             node_path = f"{path}.{safe_name}" if path else safe_name

# #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# #             current_nodes[item['id']] = {
# #                 'name': item['name'],
# #                 'parent_id': parent_id,
# #                 'type': node_type,
# #                 'mime_type': mime_type,
# #                 'path': node_path,
# #                 'modified_time': modified_time
# #             }

# #             if node_type == "folder":
# #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# #     walk_folder(root_folder_id)

# #     # --- Postgres Sync (with folder tracking) ---
# #     with conn.cursor() as cur:
# #         # Get existing nodes (both files and folders) to track changes
# #         cur.execute("SELECT id, type, modified_time FROM nodes")
# #         existing_nodes = {row[0]: {'type': row[1], 'modified_time': row[2]} for row in cur.fetchall()}
        
# #         for gid, info in current_nodes.items():
# #             is_new = gid not in existing_nodes
# #             is_updated = (not is_new and 
# #                          existing_nodes[gid]['modified_time'] != info['modified_time'])
            
# #             # Track changes by type
# #             if is_new:
# #                 if info['type'] == 'file':
# #                     stats['new_files'].append(info['name'])
# #                 else:
# #                     stats['new_folders'].append(info['name'])
# #             elif is_updated:
# #                 if info['type'] == 'file':
# #                     stats['updated_files'].append(info['name'])
# #                 else:
# #                     stats['updated_folders'].append(info['name'])
# #             elif info['type'] == 'file':
# #                 stats['unchanged_files'].append(info['name'])
            
# #             cur.execute("""
# #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# #                 ON CONFLICT (id) DO UPDATE
# #                 SET name = EXCLUDED.name,
# #                     parent_id = EXCLUDED.parent_id,
# #                     type = EXCLUDED.type,
# #                     mime_type = EXCLUDED.mime_type,
# #                     path = EXCLUDED.path,
# #                     modified_time = EXCLUDED.modified_time
# #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# #             """, (gid, info['name'], info['parent_id'], info['type'],
# #                   info['mime_type'], info['path'], info['modified_time']))

# #         # Track deleted nodes
# #         if current_nodes:
# #             deleted_ids = set(existing_nodes.keys()) - set(current_nodes.keys())
# #             if deleted_ids:
# #                 cur.execute("SELECT name, type FROM nodes WHERE id IN %s", (tuple(deleted_ids),))
# #                 for name, node_type in cur.fetchall():
# #                     if node_type == 'file':
# #                         stats['deleted_files'].append(name)
# #                     else:
# #                         stats['deleted_folders'].append(name)
            
# #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# #         else:
# #             cur.execute("SELECT name, type FROM nodes")
# #             for name, node_type in cur.fetchall():
# #                 if node_type == 'file':
# #                     stats['deleted_files'].append(name)
# #                 else:
# #                     stats['deleted_folders'].append(name)
# #             cur.execute("DELETE FROM nodes;")
# #         conn.commit()

# #     # --- Qdrant Sync (INCREMENTAL - files only) ---
# #     ensure_qdrant_collection()
    
# #     try:
# #         existing_points = qdrant.scroll(
# #             collection_name=COLLECTION_NAME,
# #             limit=10000,
# #             with_payload=True,
# #             with_vectors=False
# #         )[0]
# #         existing_drive_ids = {point.payload.get('drive_id'): {
# #             'modified_time': point.payload.get('modified_time'),
# #             'point_id': point.id
# #         } for point in existing_points}
# #     except Exception as e:
# #         print(f"⚠️ Could not fetch existing Qdrant points: {e}")
# #         existing_drive_ids = {}
    
# #     files_to_add = []
# #     files_to_update = []
# #     files_to_delete = []
    
# #     # Find files that need to be added or updated in Qdrant
# #     for gid, info in current_nodes.items():
# #         if info['type'] != "file" or not info['mime_type']:
# #             continue
            
# #         is_new = gid not in existing_drive_ids
# #         is_updated = (not is_new and 
# #                      existing_drive_ids[gid]['modified_time'] != info['modified_time'].isoformat())
        
# #         if is_new:
# #             files_to_add.append((gid, info))
# #         elif is_updated:
# #             files_to_update.append((gid, info))
    
# #     # Find files in Qdrant that have been deleted from Drive
# #     current_drive_file_ids = {gid for gid, info in current_nodes.items() if info['type'] == 'file'}
# #     for drive_id, point_info in existing_drive_ids.items():
# #         if drive_id not in current_drive_file_ids:
# #             files_to_delete.append((drive_id, point_info['point_id']))
    
# #     points_to_upsert = []
# #     qdrant_stats = {
# #         'new': [],
# #         'updated': [],
# #         'deleted': [],
# #         'skipped': []
# #     }
    
# #     if files_to_add or files_to_update:
# #         print(f"\n🔄 Processing {len(files_to_add)} new files and {len(files_to_update)} updated files for Qdrant...")
    
# #     for gid, info in files_to_add + files_to_update:
# #         is_new_file = (gid, info) in files_to_add
        
# #         text_content = extract_file_content(service, gid, info['mime_type'], info['name'])
        
# #         if text_content is None:
# #             stats['skipped_files'].append(info['name'])
# #             continue
        
# #         if is_new_file:
# #             qdrant_stats['new'].append(info['name'])
# #         else:
# #             qdrant_stats['updated'].append(info['name'])
        
# #         point_id = drive_id_to_uuid(gid)
# #         content_to_embed = text_content[:8000]
        
# #         try:
# #             embedding = generate_embedding(content_to_embed)
# #         except Exception as e:
# #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# #             continue
        
# #         points_to_upsert.append(qmodels.PointStruct(
# #             id=point_id,
# #             vector=embedding,
# #             payload={
# #                 "drive_id": gid,
# #                 "name": info['name'],
# #                 "type": info['type'],
# #                 "mime_type": info['mime_type'],
# #                 "path": info['path'],
# #                 "modified_time": info['modified_time'].isoformat(),
# #                 "content": text_content[:10000]
# #             }
# #         ))
    
# #     # Upload new/updated points to Qdrant in batches
# #     if points_to_upsert:
# #         total_points = len(points_to_upsert)
# #         print(f"\n📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# #         for i in range(0, total_points, BATCH_SIZE):
# #             batch = points_to_upsert[i:i + BATCH_SIZE]
# #             try:
# #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# #             except Exception as e:
# #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# #                 continue
        
# #         print(f"✅ Successfully synced {total_points} items to Qdrant")
    
# #     # Delete removed files from Qdrant
# #     if files_to_delete:
# #         print(f"\n🗑️  Deleting {len(files_to_delete)} removed files from Qdrant...")
# #         point_ids_to_delete = [point_id for _, point_id in files_to_delete]
        
# #         try:
# #             qdrant.delete(
# #                 collection_name=COLLECTION_NAME,
# #                 points_selector=qmodels.PointIdsList(points=point_ids_to_delete)
# #             )
# #             qdrant_stats['deleted'] = [name for name in stats['deleted_files']]
# #             print(f"✅ Deleted {len(point_ids_to_delete)} points from Qdrant")
# #         except Exception as e:
# #             print(f"❌ Failed to delete points from Qdrant: {e}")
    
# #     # Print detailed statistics
# #     print("\n" + "="*60)
# #     print("📊 INCREMENTAL SYNC SUMMARY")
# #     print("="*60)
    
# #     print("\n🗄️  PostgreSQL Changes:")
    
# #     # Files
# #     if stats['new_files']:
# #         print(f"  ✨ New files ({len(stats['new_files'])}):")
# #         for name in stats['new_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['new_files']) > 10:
# #             print(f"     ... and {len(stats['new_files']) - 10} more")
    
# #     if stats['updated_files']:
# #         print(f"  🔄 Updated files ({len(stats['updated_files'])}):")
# #         for name in stats['updated_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['updated_files']) > 10:
# #             print(f"     ... and {len(stats['updated_files']) - 10} more")
    
# #     if stats['deleted_files']:
# #         print(f"  🗑️  Deleted files ({len(stats['deleted_files'])}):")
# #         for name in stats['deleted_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['deleted_files']) > 10:
# #             print(f"     ... and {len(stats['deleted_files']) - 10} more")
    
# #     if stats['unchanged_files']:
# #         print(f"  ✓ Unchanged files: {len(stats['unchanged_files'])} (not processed)")
    
# #     # Folders
# #     if stats['new_folders']:
# #         print(f"\n  📁 New folders ({len(stats['new_folders'])}):")
# #         for name in stats['new_folders'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['new_folders']) > 10:
# #             print(f"     ... and {len(stats['new_folders']) - 10} more")
    
# #     if stats['updated_folders']:
# #         print(f"  📁 Updated folders ({len(stats['updated_folders'])}):")
# #         for name in stats['updated_folders'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['updated_folders']) > 10:
# #             print(f"     ... and {len(stats['updated_folders']) - 10} more")
    
# #     if stats['deleted_folders']:
# #         print(f"  📁 Deleted folders ({len(stats['deleted_folders'])}):")
# #         for name in stats['deleted_folders'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['deleted_folders']) > 10:
# #             print(f"     ... and {len(stats['deleted_folders']) - 10} more")
    
# #     print("\n🔍 Qdrant Vector Database Changes (Files Only):")
# #     if qdrant_stats['new']:
# #         print(f"  ✨ New embeddings ({len(qdrant_stats['new'])}):")
# #         for name in qdrant_stats['new'][:10]:
# #             print(f"     • {name}")
# #         if len(qdrant_stats['new']) > 10:
# #             print(f"     ... and {len(qdrant_stats['new']) - 10} more")
    
# #     if qdrant_stats['updated']:
# #         print(f"  🔄 Updated embeddings ({len(qdrant_stats['updated'])}):")
# #         for name in qdrant_stats['updated'][:10]:
# #             print(f"     • {name}")
# #         if len(qdrant_stats['updated']) > 10:
# #             print(f"     ... and {len(qdrant_stats['updated']) - 10} more")
    
# #     if qdrant_stats['deleted']:
# #         print(f"  🗑️  Deleted embeddings ({len(qdrant_stats['deleted'])}):")
# #         for name in qdrant_stats['deleted'][:10]:
# #             print(f"     • {name}")
# #         if len(qdrant_stats['deleted']) > 10:
# #             print(f"     ... and {len(qdrant_stats['deleted']) - 10} more")
    
# #     if stats['skipped_files']:
# #         print(f"\n⚠️  Skipped files ({len(stats['skipped_files'])}):")
# #         for name in stats['skipped_files'][:10]:
# #             print(f"     • {name}")
# #         if len(stats['skipped_files']) > 10:
# #             print(f"     ... and {len(stats['skipped_files']) - 10} more")
    
# #     # Summary
# #     total_pg_changes = (len(stats['new_files']) + len(stats['updated_files']) + 
# #                        len(stats['deleted_files']) + len(stats['new_folders']) + 
# #                        len(stats['updated_folders']) + len(stats['deleted_folders']))
# #     total_qdrant_changes = (len(qdrant_stats['new']) + len(qdrant_stats['updated']) + 
# #                            len(qdrant_stats['deleted']))
    
# #     if total_pg_changes == 0 and total_qdrant_changes == 0:
# #         print(f"\n  ✓ No changes detected - everything is up to date!")
    
# #     print("\n" + "="*60)


# # if __name__ == '__main__':
# #     service = authenticate()
# #     conn = psycopg2.connect(config.DATABASE_URL)

# #     print("Syncing Google Drive → Postgres + Qdrant (INCREMENTAL MODE)...")
# #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# #     print("\nDone! Only changed files and folders were processed.")

# #     conn.close()

# from __future__ import print_function
# import os
# import io
# import psycopg2
# from urllib.parse import quote_plus
# from google.oauth2.credentials import Credentials
# from google_auth_oauthlib.flow import InstalledAppFlow
# from google.auth.transport.requests import Request
# from googleapiclient.discovery import build
# from datetime import datetime
# import uuid
# import time

# # Embedding
# from openai import OpenAI

# # PDF parsing
# from PyPDF2 import PdfReader

# # Document parsing
# from docx import Document
# import pytesseract
# from PIL import Image

# # Configuration
# import config

# # Google Drive scopes
# SCOPES = [
#     'https://www.googleapis.com/auth/drive.readonly',
#     'https://www.googleapis.com/auth/drive.file'
# ]

# # Initialize OpenAI
# openai_client = OpenAI(api_key=config.OPENAI_API_KEY)

# # Embedding configuration
# EMBEDDING_DIM = 1536
# MAX_RETRIES = 3
# RETRY_DELAY = 2  # seconds


# # ----------------- Database Connection -----------------
# def get_db_connection():
#     """Create a new database connection with proper configuration."""
#     try:
#         conn = psycopg2.connect(
#             config.DATABASE_URL,
#             connect_timeout=30,
#             keepalives=1,
#             keepalives_idle=30,
#             keepalives_interval=10,
#             keepalives_count=5
#         )
#         # Set statement timeout to 5 minutes
#         with conn.cursor() as cur:
#             cur.execute("SET statement_timeout = '300000'")  # 5 minutes in ms
#         conn.commit()
#         return conn
#     except Exception as e:
#         print(f"Database connection failed: {e}")
#         raise


# def ensure_connection(conn):
#     """Check if connection is alive, reconnect if needed."""
#     try:
#         with conn.cursor() as cur:
#             cur.execute("SELECT 1")
#         return conn
#     except (psycopg2.OperationalError, psycopg2.InterfaceError):
#         print("Connection lost, reconnecting...")
#         return get_db_connection()


# # ----------------- Authentication -----------------
# def authenticate():
#     creds = None
#     if os.path.exists(config.TOKEN_FILE):
#         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
#         print("Loaded existing token.json")

#     if not creds or not creds.valid:
#         if creds and creds.expired and creds.refresh_token:
#             creds.refresh(Request())
#             print("Token refreshed")
#         else:
#             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
#             creds = flow.run_local_server(port=0)

#         with open(config.TOKEN_FILE, 'w') as token_file:
#             token_file.write(creds.to_json())
#         print("token.json saved successfully")

#     return build('drive', 'v3', credentials=creds)


# # ----------------- Embedding Helpers -----------------
# def generate_embedding(text: str):
#     """Generate embedding for text using OpenAI."""
#     response = openai_client.embeddings.create(
#         model="text-embedding-3-small",
#         input=text
#     )
#     return response.data[0].embedding


# def get_file_uuid(drive_id: str) -> str:
#     """Generate a stable UUID for the file based on Drive ID."""
#     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # ----------------- File Extraction -----------------
# def extract_file_content(service, file_id, mime_type, file_name=""):
#     """Download and extract text content depending on file type."""
#     try:
#         # Google Docs
#         if mime_type == "application/vnd.google-apps.document":
#             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
#             return doc.decode("utf-8")

#         # Google Sheets
#         elif mime_type == "application/vnd.google-apps.spreadsheet":
#             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
#             return sheet.decode("utf-8")

#         # PDF files
#         elif mime_type == "application/pdf":
#             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
#             reader = PdfReader(fh)
#             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
#             return text if text.strip() else None

#         # Word Documents (.docx)
#         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
#             doc = Document(fh)
#             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#             return text if text.strip() else None

#         # Images (JPEG, PNG, etc.) - use OCR
#         elif mime_type.startswith("image/"):
#             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
#             image = Image.open(fh)
#             text = pytesseract.image_to_string(image)
#             return text.strip() if text.strip() else None

#         # JSON files
#         elif mime_type == "application/json":
#             text = service.files().get_media(fileId=file_id).execute()
#             content = text.decode("utf-8")
#             return content if content.strip() else None

#         # Plain text files
#         elif mime_type.startswith("text/"):
#             text = service.files().get_media(fileId=file_id).execute()
#             content = text.decode("utf-8")
#             return content if content.strip() else None

#         else:
#             print(f"Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
#             return None

#     except Exception as e:
#         print(f"Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
#         return None


# # ----------------- PGVector Helpers -----------------
# def upsert_embedding(conn, gid, info, text_content):
#     """Insert or update a single embedding per file in pgvector table."""
#     file_uuid = get_file_uuid(gid)
    
#     # Use first 8000 characters for embedding (OpenAI's token limit consideration)
#     content_to_embed = text_content[:8000]
    
#     print(f"  Processing embedding for '{info['name']}'...")

#     for attempt in range(MAX_RETRIES):
#         try:
#             conn = ensure_connection(conn)
            
#             # Generate embedding
#             try:
#                 embedding = generate_embedding(content_to_embed)
#             except Exception as e:
#                 print(f"  Failed to generate embedding: {e}")
#                 raise
            
#             with conn.cursor() as cur:
#                 # Single INSERT with ON CONFLICT on drive_id
#                 cur.execute("""
#                     INSERT INTO file_embeddings 
#                     (id, drive_id, name, type, mime_type, path, modified_time, content, embedding)
#                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
#                     ON CONFLICT (drive_id) DO UPDATE
#                     SET id = EXCLUDED.id,
#                         name = EXCLUDED.name,
#                         type = EXCLUDED.type,
#                         mime_type = EXCLUDED.mime_type,
#                         path = EXCLUDED.path,
#                         modified_time = EXCLUDED.modified_time,
#                         content = EXCLUDED.content,
#                         embedding = EXCLUDED.embedding
#                 """, (
#                     file_uuid,
#                     gid,
#                     info['name'],
#                     info['type'],
#                     info['mime_type'],
#                     info['path'],
#                     info['modified_time'],
#                     text_content[:10000],  # Store more content for reference
#                     embedding
#                 ))

#                 # Update nodes table with vector reference
#                 cur.execute("UPDATE nodes SET vector_id = %s WHERE id = %s", (file_uuid, gid))

#             conn.commit()
#             print(f"  Successfully embedded 1 chunk")
#             return conn
            
#         except (psycopg2.OperationalError, psycopg2.InterfaceError) as e:
#             print(f"  Database error (attempt {attempt + 1}/{MAX_RETRIES}): {e}")
#             if attempt < MAX_RETRIES - 1:
#                 time.sleep(RETRY_DELAY * (attempt + 1))
#                 conn = get_db_connection()
#             else:
#                 print(f"  Failed to embed '{info['name']}' after {MAX_RETRIES} attempts")
#                 raise
#         except Exception as e:
#             print(f"  Unexpected error embedding '{info['name']}': {e}")
#             conn.rollback()
#             raise

#     return conn


# def delete_embedding(conn, drive_ids):
#     """Remove embeddings for deleted files and reset vector_id in nodes."""
#     try:
#         conn = ensure_connection(conn)
#         with conn.cursor() as cur:
#             cur.execute("DELETE FROM file_embeddings WHERE drive_id = ANY(%s)", (drive_ids,))
#             cur.execute("UPDATE nodes SET vector_id = NULL WHERE id = ANY(%s)", (drive_ids,))
#         conn.commit()
#         return conn
#     except Exception as e:
#         print(f"Error deleting embeddings: {e}")
#         conn.rollback()
#         raise


# # ----------------- Google Drive Sync -----------------
# def sync_drive(service, conn, root_folder_id):
#     current_nodes = {}
#     stats = {
#         'new_files': [], 'updated_files': [], 'unchanged_files': [], 'deleted_files': [],
#         'new_folders': [], 'updated_folders': [], 'deleted_folders': [], 'skipped_files': []
#     }

#     def walk_folder(folder_id, parent_id=None, path=""):
#         query = f"'{folder_id}' in parents and trashed=false"
#         results = service.files().list(
#             q=query,
#             spaces='drive',
#             includeItemsFromAllDrives=True,
#             supportsAllDrives=True,
#             fields="files(id, name, mimeType, modifiedTime)"
#         ).execute()
#         items = results.get('files', [])

#         for item in items:
#             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
#             mime_type = None if node_type == "folder" else item['mimeType']
#             safe_name = quote_plus(item['name'].replace(' ', '_'))
#             node_path = f"{path}.{safe_name}" if path else safe_name

#             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

#             current_nodes[item['id']] = {
#                 'name': item['name'],
#                 'parent_id': parent_id,
#                 'type': node_type,
#                 'mime_type': mime_type,
#                 'path': node_path,
#                 'modified_time': modified_time
#             }

#             if node_type == "folder":
#                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

#     walk_folder(root_folder_id)

#     # --- Postgres Sync for nodes ---
#     conn = ensure_connection(conn)
#     with conn.cursor() as cur:
#         cur.execute("SELECT id, type, modified_time FROM nodes")
#         existing_nodes = {row[0]: {'type': row[1], 'modified_time': row[2]} for row in cur.fetchall()}

#         for gid, info in current_nodes.items():
#             is_new = gid not in existing_nodes
#             is_updated = (not is_new and existing_nodes[gid]['modified_time'] != info['modified_time'])

#             if is_new:
#                 stats['new_files' if info['type'] == 'file' else 'new_folders'].append(info['name'])
#             elif is_updated:
#                 stats['updated_files' if info['type'] == 'file' else 'updated_folders'].append(info['name'])
#             elif info['type'] == 'file':
#                 stats['unchanged_files'].append(info['name'])

#             cur.execute("""
#                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
#                 VALUES (%s, %s, %s, %s, %s, %s, %s)
#                 ON CONFLICT (id) DO UPDATE
#                 SET name = EXCLUDED.name,
#                     parent_id = EXCLUDED.parent_id,
#                     type = EXCLUDED.type,
#                     mime_type = EXCLUDED.mime_type,
#                     path = EXCLUDED.path,
#                     modified_time = EXCLUDED.modified_time
#                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
#             """, (gid, info['name'], info['parent_id'], info['type'],
#                   info['mime_type'], info['path'], info['modified_time']))

#         # Detect deletions
#         deleted_ids = set(existing_nodes.keys()) - set(current_nodes.keys())
#         if deleted_ids:
#             cur.execute("SELECT id, name, type FROM nodes WHERE id = ANY(%s)", (list(deleted_ids),))
#             for node_id, name, node_type in cur.fetchall():
#                 if node_type == 'file':
#                     stats['deleted_files'].append(name)
#                 else:
#                     stats['deleted_folders'].append(name)
#             cur.execute("DELETE FROM nodes WHERE id = ANY(%s)", (list(deleted_ids),))

#         conn.commit()

#     # --- Get existing embeddings to determine what needs processing ---
#     with conn.cursor() as cur:
#         cur.execute("SELECT drive_id, modified_time FROM file_embeddings")
#         existing_embeddings = {row[0]: row[1] for row in cur.fetchall()}
    
#     # --- Determine which files actually need embedding ---
#     files_to_process = []
#     embedding_stats = {
#         'new_embeddings': [],
#         'updated_embeddings': [],
#         'unchanged_embeddings': []
#     }
    
#     for gid, info in current_nodes.items():
#         if info['type'] != 'file':
#             continue
        
#         # Check if file needs embedding
#         if gid not in existing_embeddings:
#             # New file - needs embedding
#             files_to_process.append((gid, info))
#             embedding_stats['new_embeddings'].append(info['name'])
#         elif existing_embeddings[gid] != info['modified_time']:
#             # File was modified - needs re-embedding
#             files_to_process.append((gid, info))
#             embedding_stats['updated_embeddings'].append(info['name'])
#         else:
#             # File unchanged - skip embedding
#             embedding_stats['unchanged_embeddings'].append(info['name'])
    
#     if not files_to_process:
#         print("\nNo files need embedding - all files are up to date!")
#     else:
#         print(f"\nProcessing {len(files_to_process)} files for embedding...")
#         print(f"  - New: {len(embedding_stats['new_embeddings'])}")
#         print(f"  - Updated: {len(embedding_stats['updated_embeddings'])}")
#         print(f"  - Unchanged (skipped): {len(embedding_stats['unchanged_embeddings'])}")
    
#     print(f"\nProcessing {len(files_to_process)} files for embedding...")
#     for idx, (gid, info) in enumerate(files_to_process, 1):
#         print(f"\n[{idx}/{len(files_to_process)}] Processing: {info['name']}")
#         text_content = extract_file_content(service, gid, info['mime_type'], info['name'])
#         if not text_content:
#             stats['skipped_files'].append(info['name'])
#             print(f"  Skipped (no content extracted)")
#             continue
        
#         try:
#             conn = upsert_embedding(conn, gid, info, text_content)
#         except Exception as e:
#             print(f"  Failed to process '{info['name']}': {e}")
#             stats['skipped_files'].append(info['name'])

#     # --- Delete embeddings for removed files ---
#     if deleted_ids:
#         file_deleted_ids = [gid for gid in deleted_ids if gid in existing_nodes and existing_nodes[gid]['type'] == 'file']
#         if file_deleted_ids:
#             conn = delete_embedding(conn, file_deleted_ids)

#     print("\n" + "="*60)
#     print("SYNC SUMMARY")
#     print("="*60)
    
#     print("\nPostgreSQL Metadata Changes:")
#     print(f"  New files: {len(stats['new_files'])}")
#     print(f"  Updated files: {len(stats['updated_files'])}")
#     print(f"  Deleted files: {len(stats['deleted_files'])}")
#     print(f"  New folders: {len(stats['new_folders'])}")
#     print(f"  Updated folders: {len(stats['updated_folders'])}")
#     print(f"  Deleted folders: {len(stats['deleted_folders'])}")
    
#     print("\nVector Embeddings (Actual Work Done):")
#     print(f"  New embeddings created: {len(embedding_stats['new_embeddings'])}")
#     print(f"  Embeddings updated: {len(embedding_stats['updated_embeddings'])}")
#     print(f"  Unchanged (skipped): {len(embedding_stats['unchanged_embeddings'])}")
#     print(f"  Skipped (no content): {len(stats['skipped_files'])}")
    
#     if embedding_stats['new_embeddings']:
#         print(f"\n  New files embedded:")
#         for name in embedding_stats['new_embeddings'][:5]:
#             print(f"    - {name}")
#         if len(embedding_stats['new_embeddings']) > 5:
#             print(f"    ... and {len(embedding_stats['new_embeddings']) - 5} more")
    
#     if embedding_stats['updated_embeddings']:
#         print(f"\n  Updated files re-embedded:")
#         for name in embedding_stats['updated_embeddings'][:5]:
#             print(f"    - {name}")
#         if len(embedding_stats['updated_embeddings']) > 5:
#             print(f"    ... and {len(embedding_stats['updated_embeddings']) - 5} more")
    
#     print("="*60)

#     return conn


# # ----------------- Main -----------------
# if __name__ == '__main__':
#     conn = None
#     try:
#         service = authenticate()
#         conn = get_db_connection()

#         print("Syncing Google Drive → Supabase Postgres (pgvector)...")
#         conn = sync_drive(service, conn, config.ROOT_FOLDER_ID)
#         print("\nDone! Only changed files and folders were processed.")

#     except KeyboardInterrupt:
#         print("\nProcess interrupted by user")
#     except Exception as e:
#         print(f"\nFatal error: {e}")
#         import traceback
#         traceback.print_exc()
#     finally:
#         if conn:
#             conn.close()
#             print("Database connection closed")

# # # # # from __future__ import print_function
# # # # # import os
# # # # # import io
# # # # # import psycopg2
# # # # # from urllib.parse import quote_plus
# # # # # from google.oauth2.credentials import Credentials
# # # # # from google_auth_oauthlib.flow import InstalledAppFlow
# # # # # from google.auth.transport.requests import Request
# # # # # from googleapiclient.discovery import build
# # # # # from datetime import datetime
# # # # # import uuid

# # # # # # Embedding + Vector DB
# # # # # from openai import OpenAI
# # # # # from qdrant_client import QdrantClient
# # # # # from qdrant_client.http import models as qmodels

# # # # # # PDF parsing
# # # # # from PyPDF2 import PdfReader

# # # # # # Configuration
# # # # # import config

# # # # # # Google Drive scopes (read & write)
# # # # # SCOPES = [
# # # # #     'https://www.googleapis.com/auth/drive.readonly',  # read all files
# # # # #     'https://www.googleapis.com/auth/drive.file'       # read/write files created or opened by app
# # # # # ]

# # # # # # Initialize OpenAI and Qdrant
# # # # # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # # # # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)  # Increased timeout
# # # # # COLLECTION_NAME = "fileData"  # your Qdrant collection name

# # # # # # Batch size for Qdrant upsert
# # # # # BATCH_SIZE = 100  # Process 100 points at a time


# # # # # def authenticate():
# # # # #     creds = None
# # # # #     if os.path.exists(config.TOKEN_FILE):
# # # # #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# # # # #         print("Loaded existing token.json")

# # # # #     if not creds or not creds.valid:
# # # # #         if creds and creds.expired and creds.refresh_token:
# # # # #             creds.refresh(Request())
# # # # #             print("Token refreshed")
# # # # #         else:
# # # # #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# # # # #             creds = flow.run_local_server(port=0)

# # # # #         with open(config.TOKEN_FILE, 'w') as token_file:
# # # # #             token_file.write(creds.to_json())
# # # # #         print("token.json saved successfully")

# # # # #     return build('drive', 'v3', credentials=creds)


# # # # # def ensure_qdrant_collection():
# # # # #     if not qdrant.collection_exists(COLLECTION_NAME):
# # # # #         qdrant.create_collection(
# # # # #             collection_name=COLLECTION_NAME,
# # # # #             vectors_config=qmodels.VectorParams(
# # # # #                 size=1536,
# # # # #                 distance="Cosine"
# # # # #             )
# # # # #         )
# # # # #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # # # # def generate_embedding(text: str):
# # # # #     response = openai_client.embeddings.create(
# # # # #         model="text-embedding-3-small",
# # # # #         input=text
# # # # #     )
# # # # #     return response.data[0].embedding


# # # # # def drive_id_to_uuid(drive_id: str) -> str:
# # # # #     """Convert Google Drive ID to a valid UUID for Qdrant."""
# # # # #     # Use UUID v5 with a namespace to create deterministic UUIDs
# # # # #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # # # def extract_file_content(service, file_id, mime_type, file_name=""):
# # # # #     """Download and extract text content depending on file type."""
# # # # #     try:
# # # # #         if mime_type == "application/vnd.google-apps.document":
# # # # #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# # # # #             return doc.decode("utf-8")

# # # # #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# # # # #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# # # # #             return sheet.decode("utf-8")

# # # # #         elif mime_type == "application/pdf":
# # # # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # # # #             reader = PdfReader(fh)
# # # # #             return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

# # # # #         elif mime_type.startswith("text/"):
# # # # #             text = service.files().get_media(fileId=file_id).execute()
# # # # #             content = text.decode("utf-8")
# # # # #             # Check if file is empty
# # # # #             if not content or content.strip() == "":
# # # # #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# # # # #                 return None
# # # # #             return content

# # # # #         else:
# # # # #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# # # # #             return None
# # # # #     except Exception as e:
# # # # #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# # # # #         return None


# # # # # def sync_drive(service, conn, root_folder_id):
# # # # #     current_nodes = {}

# # # # #     def walk_folder(folder_id, parent_id=None, path=""):
# # # # #         query = f"'{folder_id}' in parents and trashed=false"
# # # # #         results = service.files().list(
# # # # #             q=query,
# # # # #             spaces='drive',
# # # # #             includeItemsFromAllDrives=True,
# # # # #             supportsAllDrives=True,
# # # # #             fields="files(id, name, mimeType, modifiedTime)"
# # # # #         ).execute()
# # # # #         items = results.get('files', [])

# # # # #         for item in items:
# # # # #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# # # # #             mime_type = None if node_type == "folder" else item['mimeType']
# # # # #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# # # # #             node_path = f"{path}.{safe_name}" if path else safe_name

# # # # #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# # # # #             current_nodes[item['id']] = {
# # # # #                 'name': item['name'],
# # # # #                 'parent_id': parent_id,
# # # # #                 'type': node_type,
# # # # #                 'mime_type': mime_type,
# # # # #                 'path': node_path,
# # # # #                 'modified_time': modified_time
# # # # #             }

# # # # #             if node_type == "folder":
# # # # #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# # # # #     walk_folder(root_folder_id)

# # # # #     # --- Postgres Sync ---
# # # # #     with conn.cursor() as cur:
# # # # #         for gid, info in current_nodes.items():
# # # # #             cur.execute("""
# # # # #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# # # # #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# # # # #                 ON CONFLICT (id) DO UPDATE
# # # # #                 SET name = EXCLUDED.name,
# # # # #                     parent_id = EXCLUDED.parent_id,
# # # # #                     type = EXCLUDED.type,
# # # # #                     mime_type = EXCLUDED.mime_type,
# # # # #                     path = EXCLUDED.path,
# # # # #                     modified_time = EXCLUDED.modified_time
# # # # #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# # # # #             """, (gid, info['name'], info['parent_id'], info['type'],
# # # # #                   info['mime_type'], info['path'], info['modified_time']))

# # # # #         if current_nodes:
# # # # #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# # # # #         else:
# # # # #             cur.execute("DELETE FROM nodes;")
# # # # #         conn.commit()

# # # # #     # --- Qdrant Sync ---
# # # # #     ensure_qdrant_collection()
# # # # #     points = []

# # # # #     for gid, info in current_nodes.items():
# # # # #         text_content = None
# # # # #         if info['type'] == "file" and info['mime_type']:
# # # # #             text_content = extract_file_content(service, gid, info['mime_type'], info['name'])

# # # # #         if text_content is None:
# # # # #             # Skip files we couldn't read (empty, unsupported, or errored)
# # # # #             continue

# # # # #         # --- Convert Drive ID to deterministic UUID for Qdrant ---
# # # # #         point_id = drive_id_to_uuid(gid)
# # # # #         content_to_embed = text_content[:8000]  # limit to 8k chars
        
# # # # #         try:
# # # # #             embedding = generate_embedding(content_to_embed)
# # # # #         except Exception as e:
# # # # #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# # # # #             continue

# # # # #         points.append(qmodels.PointStruct(
# # # # #             id=point_id,
# # # # #             vector=embedding,
# # # # #             payload={
# # # # #                 "drive_id": gid,  # original Google Drive ID
# # # # #                 "name": info['name'],
# # # # #                 "type": info['type'],
# # # # #                 "mime_type": info['mime_type'],
# # # # #                 "path": info['path'],
# # # # #                 "modified_time": info['modified_time'].isoformat(),
# # # # #                 "content": text_content[:10000]  # Store limited content in payload
# # # # #             }
# # # # #         ))

# # # # #     # Upload to Qdrant in batches to avoid timeout
# # # # #     if points:
# # # # #         total_points = len(points)
# # # # #         print(f"📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# # # # #         for i in range(0, total_points, BATCH_SIZE):
# # # # #             batch = points[i:i + BATCH_SIZE]
# # # # #             try:
# # # # #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# # # # #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# # # # #             except Exception as e:
# # # # #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# # # # #                 # Continue with next batch instead of failing completely
# # # # #                 continue
        
# # # # #         print(f"✅ Successfully synced {total_points} items to Qdrant")
# # # # #     else:
# # # # #         print("⚠️ No points to upload to Qdrant")


# # # # # if __name__ == '__main__':
# # # # #     service = authenticate()
# # # # #     conn = psycopg2.connect(config.DATABASE_URL)

# # # # #     print("Syncing Google Drive → Postgres + Qdrant (with file contents)...")
# # # # #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# # # # #     print("Done! DB + Qdrant are now consistent with Drive.")

# # # # #     conn.close()

# # # # from __future__ import print_function
# # # # import os
# # # # import io
# # # # import psycopg2
# # # # from urllib.parse import quote_plus
# # # # from google.oauth2.credentials import Credentials
# # # # from google_auth_oauthlib.flow import InstalledAppFlow
# # # # from google.auth.transport.requests import Request
# # # # from googleapiclient.discovery import build
# # # # from datetime import datetime
# # # # import uuid

# # # # # Embedding + Vector DB
# # # # from openai import OpenAI
# # # # from qdrant_client import QdrantClient
# # # # from qdrant_client.http import models as qmodels

# # # # # PDF parsing
# # # # from PyPDF2 import PdfReader

# # # # # Document parsing
# # # # from docx import Document  # python-docx for Word documents
# # # # import pytesseract  # OCR for images
# # # # from PIL import Image  # Image processing

# # # # # Configuration
# # # # import config

# # # # # Google Drive scopes (read & write)
# # # # SCOPES = [
# # # #     'https://www.googleapis.com/auth/drive.readonly',  # read all files
# # # #     'https://www.googleapis.com/auth/drive.file'       # read/write files created or opened by app
# # # # ]

# # # # # Initialize OpenAI and Qdrant
# # # # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # # # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)  # Increased timeout
# # # # COLLECTION_NAME = "fileData"  # your Qdrant collection name

# # # # # Batch size for Qdrant upsert
# # # # BATCH_SIZE = 100  # Process 100 points at a time


# # # # def authenticate():
# # # #     creds = None
# # # #     if os.path.exists(config.TOKEN_FILE):
# # # #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# # # #         print("Loaded existing token.json")

# # # #     if not creds or not creds.valid:
# # # #         if creds and creds.expired and creds.refresh_token:
# # # #             creds.refresh(Request())
# # # #             print("Token refreshed")
# # # #         else:
# # # #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# # # #             creds = flow.run_local_server(port=0)

# # # #         with open(config.TOKEN_FILE, 'w') as token_file:
# # # #             token_file.write(creds.to_json())
# # # #         print("token.json saved successfully")

# # # #     return build('drive', 'v3', credentials=creds)


# # # # def ensure_qdrant_collection():
# # # #     if not qdrant.collection_exists(COLLECTION_NAME):
# # # #         qdrant.create_collection(
# # # #             collection_name=COLLECTION_NAME,
# # # #             vectors_config=qmodels.VectorParams(
# # # #                 size=1536,
# # # #                 distance="Cosine"
# # # #             )
# # # #         )
# # # #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # # # def generate_embedding(text: str):
# # # #     response = openai_client.embeddings.create(
# # # #         model="text-embedding-3-small",
# # # #         input=text
# # # #     )
# # # #     return response.data[0].embedding


# # # # def drive_id_to_uuid(drive_id: str) -> str:
# # # #     """Convert Google Drive ID to a valid UUID for Qdrant."""
# # # #     # Use UUID v5 with a namespace to create deterministic UUIDs
# # # #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # # def extract_file_content(service, file_id, mime_type, file_name=""):
# # # #     """Download and extract text content depending on file type."""
# # # #     try:
# # # #         # Google Docs
# # # #         if mime_type == "application/vnd.google-apps.document":
# # # #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# # # #             return doc.decode("utf-8")

# # # #         # Google Sheets
# # # #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# # # #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# # # #             return sheet.decode("utf-8")

# # # #         # PDF files
# # # #         elif mime_type == "application/pdf":
# # # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # # #             reader = PdfReader(fh)
# # # #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# # # #             if not text or text.strip() == "":
# # # #                 print(f"⚠️ PDF '{file_name}' has no extractable text")
# # # #                 return None
# # # #             return text

# # # #         # Word Documents (.docx)
# # # #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # # #             doc = Document(fh)
# # # #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# # # #             if not text or text.strip() == "":
# # # #                 print(f"⚠️ Word document '{file_name}' is empty")
# # # #                 return None
# # # #             return text

# # # #         # Images (JPEG, PNG, GIF, BMP, TIFF) - use OCR
# # # #         elif mime_type.startswith("image/"):
# # # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # # #             image = Image.open(fh)
            
# # # #             # Perform OCR
# # # #             text = pytesseract.image_to_string(image)
            
# # # #             if not text or text.strip() == "":
# # # #                 print(f"⚠️ Image '{file_name}' contains no readable text (OCR found nothing)")
# # # #                 return None
            
# # # #             return text.strip()

# # # #         # Plain text files
# # # #         elif mime_type.startswith("text/"):
# # # #             text = service.files().get_media(fileId=file_id).execute()
# # # #             content = text.decode("utf-8")
# # # #             # Check if file is empty
# # # #             if not content or content.strip() == "":
# # # #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# # # #                 return None
# # # #             return content

# # # #         else:
# # # #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# # # #             return None
            
# # # #     except Exception as e:
# # # #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# # # #         return None


# # # # def sync_drive(service, conn, root_folder_id):
# # # #     current_nodes = {}
# # # #     stats = {
# # # #         'new_files': [],
# # # #         'updated_files': [],
# # # #         'unchanged_files': [],
# # # #         'deleted_files': [],
# # # #         'skipped_files': []
# # # #     }

# # # #     def walk_folder(folder_id, parent_id=None, path=""):
# # # #         query = f"'{folder_id}' in parents and trashed=false"
# # # #         results = service.files().list(
# # # #             q=query,
# # # #             spaces='drive',
# # # #             includeItemsFromAllDrives=True,
# # # #             supportsAllDrives=True,
# # # #             fields="files(id, name, mimeType, modifiedTime)"
# # # #         ).execute()
# # # #         items = results.get('files', [])

# # # #         for item in items:
# # # #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# # # #             mime_type = None if node_type == "folder" else item['mimeType']
# # # #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# # # #             node_path = f"{path}.{safe_name}" if path else safe_name

# # # #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# # # #             current_nodes[item['id']] = {
# # # #                 'name': item['name'],
# # # #                 'parent_id': parent_id,
# # # #                 'type': node_type,
# # # #                 'mime_type': mime_type,
# # # #                 'path': node_path,
# # # #                 'modified_time': modified_time
# # # #             }

# # # #             if node_type == "folder":
# # # #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# # # #     walk_folder(root_folder_id)

# # # #     # --- Postgres Sync ---
# # # #     with conn.cursor() as cur:
# # # #         # Get existing files to track changes
# # # #         cur.execute("SELECT id, modified_time FROM nodes WHERE type = 'file'")
# # # #         existing_files = {row[0]: row[1] for row in cur.fetchall()}
        
# # # #         for gid, info in current_nodes.items():
# # # #             # Check if file is new or updated
# # # #             if gid not in existing_files:
# # # #                 if info['type'] == 'file':
# # # #                     stats['new_files'].append(info['name'])
# # # #             elif info['type'] == 'file' and existing_files[gid] != info['modified_time']:
# # # #                 stats['updated_files'].append(info['name'])
# # # #             elif info['type'] == 'file':
# # # #                 stats['unchanged_files'].append(info['name'])
            
# # # #             cur.execute("""
# # # #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# # # #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# # # #                 ON CONFLICT (id) DO UPDATE
# # # #                 SET name = EXCLUDED.name,
# # # #                     parent_id = EXCLUDED.parent_id,
# # # #                     type = EXCLUDED.type,
# # # #                     mime_type = EXCLUDED.mime_type,
# # # #                     path = EXCLUDED.path,
# # # #                     modified_time = EXCLUDED.modified_time
# # # #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# # # #             """, (gid, info['name'], info['parent_id'], info['type'],
# # # #                   info['mime_type'], info['path'], info['modified_time']))

# # # #         # Track deleted files
# # # #         if current_nodes:
# # # #             current_file_ids = [gid for gid, info in current_nodes.items() if info['type'] == 'file']
# # # #             deleted_ids = set(existing_files.keys()) - set(current_file_ids)
# # # #             if deleted_ids:
# # # #                 cur.execute("SELECT name FROM nodes WHERE id IN %s", (tuple(deleted_ids),))
# # # #                 stats['deleted_files'] = [row[0] for row in cur.fetchall()]
            
# # # #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# # # #         else:
# # # #             cur.execute("SELECT name FROM nodes WHERE type = 'file'")
# # # #             stats['deleted_files'] = [row[0] for row in cur.fetchall()]
# # # #             cur.execute("DELETE FROM nodes;")
# # # #         conn.commit()

# # # #     # --- Qdrant Sync ---
# # # #     ensure_qdrant_collection()
    
# # # #     # Get existing points from Qdrant to track what's new/updated
# # # #     try:
# # # #         existing_points = qdrant.scroll(
# # # #             collection_name=COLLECTION_NAME,
# # # #             limit=10000,
# # # #             with_payload=True,
# # # #             with_vectors=False
# # # #         )[0]
# # # #         existing_drive_ids = {point.payload.get('drive_id'): point.payload.get('modified_time') 
# # # #                               for point in existing_points}
# # # #     except Exception as e:
# # # #         print(f"⚠️ Could not fetch existing Qdrant points: {e}")
# # # #         existing_drive_ids = {}
    
# # # #     points = []
# # # #     qdrant_stats = {
# # # #         'new': [],
# # # #         'updated': [],
# # # #         'skipped': []
# # # #     }

# # # #     for gid, info in current_nodes.items():
# # # #         text_content = None
# # # #         if info['type'] == "file" and info['mime_type']:
# # # #             text_content = extract_file_content(service, gid, info['mime_type'], info['name'])

# # # #         if text_content is None:
# # # #             # Skip files we couldn't read (empty, unsupported, or errored)
# # # #             stats['skipped_files'].append(info['name'])
# # # #             continue
        
# # # #         # Check if this is new or updated in Qdrant
# # # #         is_new = gid not in existing_drive_ids
# # # #         is_updated = (not is_new and 
# # # #                      existing_drive_ids[gid] != info['modified_time'].isoformat())
        
# # # #         if is_new:
# # # #             qdrant_stats['new'].append(info['name'])
# # # #         elif is_updated:
# # # #             qdrant_stats['updated'].append(info['name'])
# # # #         else:
# # # #             # File exists and hasn't changed, but we'll re-upload to ensure consistency
# # # #             pass

# # # #         # --- Convert Drive ID to deterministic UUID for Qdrant ---
# # # #         point_id = drive_id_to_uuid(gid)
# # # #         content_to_embed = text_content[:8000]  # limit to 8k chars
        
# # # #         try:
# # # #             embedding = generate_embedding(content_to_embed)
# # # #         except Exception as e:
# # # #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# # # #             continue

# # # #         points.append(qmodels.PointStruct(
# # # #             id=point_id,
# # # #             vector=embedding,
# # # #             payload={
# # # #                 "drive_id": gid,  # original Google Drive ID
# # # #                 "name": info['name'],
# # # #                 "type": info['type'],
# # # #                 "mime_type": info['mime_type'],
# # # #                 "path": info['path'],
# # # #                 "modified_time": info['modified_time'].isoformat(),
# # # #                 "content": text_content[:10000]  # Store limited content in payload
# # # #             }
# # # #         ))

# # # #     # Upload to Qdrant in batches to avoid timeout
# # # #     if points:
# # # #         total_points = len(points)
# # # #         print(f"\n📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# # # #         for i in range(0, total_points, BATCH_SIZE):
# # # #             batch = points[i:i + BATCH_SIZE]
# # # #             try:
# # # #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# # # #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# # # #             except Exception as e:
# # # #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# # # #                 # Continue with next batch instead of failing completely
# # # #                 continue
        
# # # #         print(f"✅ Successfully synced {total_points} items to Qdrant")
# # # #     else:
# # # #         print("⚠️ No points to upload to Qdrant")
    
# # # #     # Print detailed statistics
# # # #     print("\n" + "="*60)
# # # #     print("📊 SYNC SUMMARY")
# # # #     print("="*60)
    
# # # #     print("\n🗄️  PostgreSQL Changes:")
# # # #     if stats['new_files']:
# # # #         print(f"  ✨ New files ({len(stats['new_files'])}):")
# # # #         for name in stats['new_files'][:10]:  # Show first 10
# # # #             print(f"     • {name}")
# # # #         if len(stats['new_files']) > 10:
# # # #             print(f"     ... and {len(stats['new_files']) - 10} more")
    
# # # #     if stats['updated_files']:
# # # #         print(f"  🔄 Updated files ({len(stats['updated_files'])}):")
# # # #         for name in stats['updated_files'][:10]:
# # # #             print(f"     • {name}")
# # # #         if len(stats['updated_files']) > 10:
# # # #             print(f"     ... and {len(stats['updated_files']) - 10} more")
    
# # # #     if stats['deleted_files']:
# # # #         print(f"  🗑️  Deleted files ({len(stats['deleted_files'])}):")
# # # #         for name in stats['deleted_files'][:10]:
# # # #             print(f"     • {name}")
# # # #         if len(stats['deleted_files']) > 10:
# # # #             print(f"     ... and {len(stats['deleted_files']) - 10} more")
    
# # # #     if stats['unchanged_files']:
# # # #         print(f"  ✓ Unchanged files: {len(stats['unchanged_files'])}")
    
# # # #     print("\n🔍 Qdrant Vector Database Changes:")
# # # #     if qdrant_stats['new']:
# # # #         print(f"  ✨ New embeddings ({len(qdrant_stats['new'])}):")
# # # #         for name in qdrant_stats['new'][:10]:
# # # #             print(f"     • {name}")
# # # #         if len(qdrant_stats['new']) > 10:
# # # #             print(f"     ... and {len(qdrant_stats['new']) - 10} more")
    
# # # #     if qdrant_stats['updated']:
# # # #         print(f"  🔄 Updated embeddings ({len(qdrant_stats['updated'])}):")
# # # #         for name in qdrant_stats['updated'][:10]:
# # # #             print(f"     • {name}")
# # # #         if len(qdrant_stats['updated']) > 10:
# # # #             print(f"     ... and {len(qdrant_stats['updated']) - 10} more")
    
# # # #     if stats['skipped_files']:
# # # #         print(f"\n⚠️  Skipped files ({len(stats['skipped_files'])}):")
# # # #         for name in stats['skipped_files'][:10]:
# # # #             print(f"     • {name}")
# # # #         if len(stats['skipped_files']) > 10:
# # # #             print(f"     ... and {len(stats['skipped_files']) - 10} more")
    
# # # #     print("\n" + "="*60)


# # # # if __name__ == '__main__':
# # # #     service = authenticate()
# # # #     conn = psycopg2.connect(config.DATABASE_URL)

# # # #     print("Syncing Google Drive → Postgres + Qdrant (with file contents)...")
# # # #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# # # #     print("Done! DB + Qdrant are now consistent with Drive.")

# # # #     conn.close()
# # # from __future__ import print_function
# # # import os
# # # import io
# # # import psycopg2
# # # from urllib.parse import quote_plus
# # # from google.oauth2.credentials import Credentials
# # # from google_auth_oauthlib.flow import InstalledAppFlow
# # # from google.auth.transport.requests import Request
# # # from googleapiclient.discovery import build
# # # from datetime import datetime
# # # import uuid

# # # # Embedding + Vector DB
# # # from openai import OpenAI
# # # from qdrant_client import QdrantClient
# # # from qdrant_client.http import models as qmodels

# # # # PDF parsing
# # # from PyPDF2 import PdfReader

# # # # Document parsing
# # # from docx import Document  # python-docx for Word documents
# # # import pytesseract  # OCR for images
# # # from PIL import Image  # Image processing

# # # # Configuration
# # # import config

# # # # Google Drive scopes (read & write)
# # # SCOPES = [
# # #     'https://www.googleapis.com/auth/drive.readonly',  # read all files
# # #     'https://www.googleapis.com/auth/drive.file'       # read/write files created or opened by app
# # # ]

# # # # Initialize OpenAI and Qdrant
# # # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)
# # # COLLECTION_NAME = "fileData"

# # # # Batch size for Qdrant upsert
# # # BATCH_SIZE = 100


# # # def authenticate():
# # #     creds = None
# # #     if os.path.exists(config.TOKEN_FILE):
# # #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# # #         print("Loaded existing token.json")

# # #     if not creds or not creds.valid:
# # #         if creds and creds.expired and creds.refresh_token:
# # #             creds.refresh(Request())
# # #             print("Token refreshed")
# # #         else:
# # #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# # #             creds = flow.run_local_server(port=0)

# # #         with open(config.TOKEN_FILE, 'w') as token_file:
# # #             token_file.write(creds.to_json())
# # #         print("token.json saved successfully")

# # #     return build('drive', 'v3', credentials=creds)


# # # def ensure_qdrant_collection():
# # #     if not qdrant.collection_exists(COLLECTION_NAME):
# # #         qdrant.create_collection(
# # #             collection_name=COLLECTION_NAME,
# # #             vectors_config=qmodels.VectorParams(
# # #                 size=1536,
# # #                 distance="Cosine"
# # #             )
# # #         )
# # #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # # def generate_embedding(text: str):
# # #     response = openai_client.embeddings.create(
# # #         model="text-embedding-3-small",
# # #         input=text
# # #     )
# # #     return response.data[0].embedding


# # # def drive_id_to_uuid(drive_id: str) -> str:
# # #     """Convert Google Drive ID to a valid UUID for Qdrant."""
# # #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # def extract_file_content(service, file_id, mime_type, file_name=""):
# # #     """Download and extract text content depending on file type."""
# # #     try:
# # #         # Google Docs
# # #         if mime_type == "application/vnd.google-apps.document":
# # #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# # #             return doc.decode("utf-8")

# # #         # Google Sheets
# # #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# # #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# # #             return sheet.decode("utf-8")

# # #         # PDF files
# # #         elif mime_type == "application/pdf":
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             reader = PdfReader(fh)
# # #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ PDF '{file_name}' has no extractable text")
# # #                 return None
# # #             return text

# # #         # Word Documents (.docx)
# # #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             doc = Document(fh)
# # #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ Word document '{file_name}' is empty")
# # #                 return None
# # #             return text

# # #         # Images (JPEG, PNG, GIF, BMP, TIFF) - use OCR
# # #         elif mime_type.startswith("image/"):
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             image = Image.open(fh)
            
# # #             # Perform OCR
# # #             text = pytesseract.image_to_string(image)
            
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ Image '{file_name}' contains no readable text (OCR found nothing)")
# # #                 return None
            
# # #             return text.strip()

# # #         # Plain text files
# # #         elif mime_type.startswith("text/"):
# # #             text = service.files().get_media(fileId=file_id).execute()
# # #             content = text.decode("utf-8")
# # #             # Check if file is empty
# # #             if not content or content.strip() == "":
# # #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# # #                 return None
# # #             return content

# # #         else:
# # #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# # #             return None
            
# # #     except Exception as e:
# # #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# # #         return None


# # # def sync_drive(service, conn, root_folder_id):
# # #     current_nodes = {}
# # #     stats = {
# # #         'new_files': [],
# # #         'updated_files': [],
# # #         'unchanged_files': [],
# # #         'deleted_files': [],
# # #         'skipped_files': []
# # #     }

# # #     def walk_folder(folder_id, parent_id=None, path=""):
# # #         query = f"'{folder_id}' in parents and trashed=false"
# # #         results = service.files().list(
# # #             q=query,
# # #             spaces='drive',
# # #             includeItemsFromAllDrives=True,
# # #             supportsAllDrives=True,
# # #             fields="files(id, name, mimeType, modifiedTime)"
# # #         ).execute()
# # #         items = results.get('files', [])

# # #         for item in items:
# # #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# # #             mime_type = None if node_type == "folder" else item['mimeType']
# # #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# # #             node_path = f"{path}.{safe_name}" if path else safe_name

# # #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# # #             current_nodes[item['id']] = {
# # #                 'name': item['name'],
# # #                 'parent_id': parent_id,
# # #                 'type': node_type,
# # #                 'mime_type': mime_type,
# # #                 'path': node_path,
# # #                 'modified_time': modified_time
# # #             }

# # #             if node_type == "folder":
# # #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# # #     walk_folder(root_folder_id)

# # #     # --- Postgres Sync ---
# # #     with conn.cursor() as cur:
# # #         # Get existing files to track changes
# # #         cur.execute("SELECT id, modified_time FROM nodes WHERE type = 'file'")
# # #         existing_files = {row[0]: row[1] for row in cur.fetchall()}
        
# # #         for gid, info in current_nodes.items():
# # #             # Check if file is new or updated
# # #             if gid not in existing_files:
# # #                 if info['type'] == 'file':
# # #                     stats['new_files'].append(info['name'])
# # #             elif info['type'] == 'file' and existing_files[gid] != info['modified_time']:
# # #                 stats['updated_files'].append(info['name'])
# # #             elif info['type'] == 'file':
# # #                 stats['unchanged_files'].append(info['name'])
            
# # #             cur.execute("""
# # #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# # #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# # #                 ON CONFLICT (id) DO UPDATE
# # #                 SET name = EXCLUDED.name,
# # #                     parent_id = EXCLUDED.parent_id,
# # #                     type = EXCLUDED.type,
# # #                     mime_type = EXCLUDED.mime_type,
# # #                     path = EXCLUDED.path,
# # #                     modified_time = EXCLUDED.modified_time
# # #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# # #             """, (gid, info['name'], info['parent_id'], info['type'],
# # #                   info['mime_type'], info['path'], info['modified_time']))

# # #         # Track deleted files
# # #         if current_nodes:
# # #             current_file_ids = [gid for gid, info in current_nodes.items() if info['type'] == 'file']
# # #             deleted_ids = set(existing_files.keys()) - set(current_file_ids)
# # #             if deleted_ids:
# # #                 cur.execute("SELECT name FROM nodes WHERE id IN %s", (tuple(deleted_ids),))
# # #                 stats['deleted_files'] = [row[0] for row in cur.fetchall()]
            
# # #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# # #         else:
# # #             cur.execute("SELECT name FROM nodes WHERE type = 'file'")
# # #             stats['deleted_files'] = [row[0] for row in cur.fetchall()]
# # #             cur.execute("DELETE FROM nodes;")
# # #         conn.commit()

# # #     # --- Qdrant Sync (INCREMENTAL) ---
# # #     ensure_qdrant_collection()
    
# # #     # Get existing points from Qdrant
# # #     try:
# # #         existing_points = qdrant.scroll(
# # #             collection_name=COLLECTION_NAME,
# # #             limit=10000,
# # #             with_payload=True,
# # #             with_vectors=False
# # #         )[0]
# # #         existing_drive_ids = {point.payload.get('drive_id'): {
# # #             'modified_time': point.payload.get('modified_time'),
# # #             'point_id': point.id
# # #         } for point in existing_points}
# # #     except Exception as e:
# # #         print(f"⚠️ Could not fetch existing Qdrant points: {e}")
# # #         existing_drive_ids = {}
    
# # #     # Identify files that need action
# # #     files_to_add = []  # New files
# # #     files_to_update = []  # Modified files
# # #     files_to_delete = []  # Deleted from Drive
    
# # #     # Find files in Drive that need to be added or updated
# # #     for gid, info in current_nodes.items():
# # #         if info['type'] != "file" or not info['mime_type']:
# # #             continue
            
# # #         is_new = gid not in existing_drive_ids
# # #         is_updated = (not is_new and 
# # #                      existing_drive_ids[gid]['modified_time'] != info['modified_time'].isoformat())
        
# # #         if is_new:
# # #             files_to_add.append((gid, info))
# # #         elif is_updated:
# # #             files_to_update.append((gid, info))
    
# # #     # Find files in Qdrant that have been deleted from Drive
# # #     current_drive_file_ids = {gid for gid, info in current_nodes.items() if info['type'] == 'file'}
# # #     for drive_id, point_info in existing_drive_ids.items():
# # #         if drive_id not in current_drive_file_ids:
# # #             files_to_delete.append((drive_id, point_info['point_id']))
    
# # #     # Process new and updated files
# # #     points_to_upsert = []
# # #     qdrant_stats = {
# # #         'new': [],
# # #         'updated': [],
# # #         'deleted': [],
# # #         'skipped': []
# # #     }
    
# # #     print(f"\n🔄 Processing {len(files_to_add)} new files and {len(files_to_update)} updated files...")
    
# # #     for gid, info in files_to_add + files_to_update:
# # #         is_new_file = (gid, info) in files_to_add
        
# # #         text_content = extract_file_content(service, gid, info['mime_type'], info['name'])
        
# # #         if text_content is None:
# # #             stats['skipped_files'].append(info['name'])
# # #             continue
        
# # #         if is_new_file:
# # #             qdrant_stats['new'].append(info['name'])
# # #         else:
# # #             qdrant_stats['updated'].append(info['name'])
        
# # #         point_id = drive_id_to_uuid(gid)
# # #         content_to_embed = text_content[:8000]
        
# # #         try:
# # #             embedding = generate_embedding(content_to_embed)
# # #         except Exception as e:
# # #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# # #             continue
        
# # #         points_to_upsert.append(qmodels.PointStruct(
# # #             id=point_id,
# # #             vector=embedding,
# # #             payload={
# # #                 "drive_id": gid,
# # #                 "name": info['name'],
# # #                 "type": info['type'],
# # #                 "mime_type": info['mime_type'],
# # #                 "path": info['path'],
# # #                 "modified_time": info['modified_time'].isoformat(),
# # #                 "content": text_content[:10000]
# # #             }
# # #         ))
    
# # #     # Upload new/updated points to Qdrant in batches
# # #     if points_to_upsert:
# # #         total_points = len(points_to_upsert)
# # #         print(f"\n📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# # #         for i in range(0, total_points, BATCH_SIZE):
# # #             batch = points_to_upsert[i:i + BATCH_SIZE]
# # #             try:
# # #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# # #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# # #             except Exception as e:
# # #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# # #                 continue
        
# # #         print(f"✅ Successfully synced {total_points} items to Qdrant")
    
# # #     # Delete removed files from Qdrant
# # #     if files_to_delete:
# # #         print(f"\n🗑️  Deleting {len(files_to_delete)} removed files from Qdrant...")
# # #         point_ids_to_delete = [point_id for _, point_id in files_to_delete]
        
# # #         try:
# # #             qdrant.delete(
# # #                 collection_name=COLLECTION_NAME,
# # #                 points_selector=qmodels.PointIdsList(points=point_ids_to_delete)
# # #             )
# # #             qdrant_stats['deleted'] = stats['deleted_files']
# # #             print(f"✅ Deleted {len(point_ids_to_delete)} points from Qdrant")
# # #         except Exception as e:
# # #             print(f"❌ Failed to delete points from Qdrant: {e}")
    
# # #     # Print detailed statistics
# # #     print("\n" + "="*60)
# # #     print("📊 INCREMENTAL SYNC SUMMARY")
# # #     print("="*60)
    
# # #     print("\n🗄️  PostgreSQL Changes:")
# # #     if stats['new_files']:
# # #         print(f"  ✨ New files ({len(stats['new_files'])}):")
# # #         for name in stats['new_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['new_files']) > 10:
# # #             print(f"     ... and {len(stats['new_files']) - 10} more")
    
# # #     if stats['updated_files']:
# # #         print(f"  🔄 Updated files ({len(stats['updated_files'])}):")
# # #         for name in stats['updated_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['updated_files']) > 10:
# # #             print(f"     ... and {len(stats['updated_files']) - 10} more")
    
# # #     if stats['deleted_files']:
# # #         print(f"  🗑️  Deleted files ({len(stats['deleted_files'])}):")
# # #         for name in stats['deleted_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['deleted_files']) > 10:
# # #             print(f"     ... and {len(stats['deleted_files']) - 10} more")
    
# # #     if stats['unchanged_files']:
# # #         print(f"  ✓ Unchanged files: {len(stats['unchanged_files'])} (not processed)")
    
# # #     print("\n🔍 Qdrant Vector Database Changes:")
# # #     if qdrant_stats['new']:
# # #         print(f"  ✨ New embeddings ({len(qdrant_stats['new'])}):")
# # #         for name in qdrant_stats['new'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['new']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['new']) - 10} more")
    
# # #     if qdrant_stats['updated']:
# # #         print(f"  🔄 Updated embeddings ({len(qdrant_stats['updated'])}):")
# # #         for name in qdrant_stats['updated'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['updated']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['updated']) - 10} more")
    
# # #     if qdrant_stats['deleted']:
# # #         print(f"  🗑️  Deleted embeddings ({len(qdrant_stats['deleted'])}):")
# # #         for name in qdrant_stats['deleted'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['deleted']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['deleted']) - 10} more")
    
# # #     if stats['skipped_files']:
# # #         print(f"\n⚠️  Skipped files ({len(stats['skipped_files'])}):")
# # #         for name in stats['skipped_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['skipped_files']) > 10:
# # #             print(f"     ... and {len(stats['skipped_files']) - 10} more")
    
# # #     if not (qdrant_stats['new'] or qdrant_stats['updated'] or qdrant_stats['deleted']):
# # #         print(f"  ✓ No changes - all files are up to date!")
    
# # #     print("\n" + "="*60)


# # # if __name__ == '__main__':
# # #     service = authenticate()
# # #     conn = psycopg2.connect(config.DATABASE_URL)

# # #     print("Syncing Google Drive → Postgres + Qdrant (INCREMENTAL MODE)...")
# # #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# # #     print("Done! Only changed files were processed.")

# # #     conn.close()

# # # from __future__ import print_function
# # # import os
# # # import io
# # # import psycopg2
# # # from urllib.parse import quote_plus
# # # from google.oauth2.credentials import Credentials
# # # from google_auth_oauthlib.flow import InstalledAppFlow
# # # from google.auth.transport.requests import Request
# # # from googleapiclient.discovery import build
# # # from datetime import datetime
# # # import uuid

# # # # Embedding + Vector DB
# # # from openai import OpenAI
# # # from qdrant_client import QdrantClient
# # # from qdrant_client.http import models as qmodels

# # # # PDF parsing
# # # from PyPDF2 import PdfReader

# # # # Document parsing
# # # from docx import Document
# # # import pytesseract
# # # from PIL import Image

# # # # Configuration
# # # import config

# # # # Google Drive scopes
# # # SCOPES = [
# # #     'https://www.googleapis.com/auth/drive.readonly',
# # #     'https://www.googleapis.com/auth/drive.file'
# # # ]

# # # # Initialize OpenAI and Qdrant
# # # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
# # # qdrant = QdrantClient(url=config.QDRANT_URL, api_key=config.QDRANT_API_KEY, timeout=120)
# # # COLLECTION_NAME = "fileData"

# # # BATCH_SIZE = 100


# # # def authenticate():
# # #     creds = None
# # #     if os.path.exists(config.TOKEN_FILE):
# # #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# # #         print("Loaded existing token.json")

# # #     if not creds or not creds.valid:
# # #         if creds and creds.expired and creds.refresh_token:
# # #             creds.refresh(Request())
# # #             print("Token refreshed")
# # #         else:
# # #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# # #             creds = flow.run_local_server(port=0)

# # #         with open(config.TOKEN_FILE, 'w') as token_file:
# # #             token_file.write(creds.to_json())
# # #         print("token.json saved successfully")

# # #     return build('drive', 'v3', credentials=creds)


# # # def ensure_qdrant_collection():
# # #     if not qdrant.collection_exists(COLLECTION_NAME):
# # #         qdrant.create_collection(
# # #             collection_name=COLLECTION_NAME,
# # #             vectors_config=qmodels.VectorParams(
# # #                 size=1536,
# # #                 distance="Cosine"
# # #             )
# # #         )
# # #         print(f"✅ Created Qdrant collection: {COLLECTION_NAME}")


# # # def generate_embedding(text: str):
# # #     response = openai_client.embeddings.create(
# # #         model="text-embedding-3-small",
# # #         input=text
# # #     )
# # #     return response.data[0].embedding


# # # def drive_id_to_uuid(drive_id: str) -> str:
# # #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # def extract_file_content(service, file_id, mime_type, file_name=""):
# # #     """Download and extract text content depending on file type."""
# # #     try:
# # #         # Google Docs
# # #         if mime_type == "application/vnd.google-apps.document":
# # #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# # #             return doc.decode("utf-8")

# # #         # Google Sheets
# # #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# # #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# # #             return sheet.decode("utf-8")

# # #         # PDF files
# # #         elif mime_type == "application/pdf":
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             reader = PdfReader(fh)
# # #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ PDF '{file_name}' has no extractable text")
# # #                 return None
# # #             return text

# # #         # Word Documents (.docx)
# # #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             doc = Document(fh)
# # #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ Word document '{file_name}' is empty")
# # #                 return None
# # #             return text

# # #         # Images (JPEG, PNG, GIF, BMP, TIFF) - use OCR
# # #         elif mime_type.startswith("image/"):
# # #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# # #             image = Image.open(fh)
            
# # #             text = pytesseract.image_to_string(image)
            
# # #             if not text or text.strip() == "":
# # #                 print(f"⚠️ Image '{file_name}' contains no readable text (OCR found nothing)")
# # #                 return None
            
# # #             return text.strip()

# # #         # Plain text files
# # #         elif mime_type.startswith("text/"):
# # #             text = service.files().get_media(fileId=file_id).execute()
# # #             content = text.decode("utf-8")
# # #             if not content or content.strip() == "":
# # #                 print(f"⚠️ File '{file_name}' ({file_id}) is empty")
# # #                 return None
# # #             return content

# # #         else:
# # #             print(f"⚠️ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# # #             return None
            
# # #     except Exception as e:
# # #         print(f"⚠️ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# # #         return None


# # # def sync_drive(service, conn, root_folder_id):
# # #     current_nodes = {}
# # #     stats = {
# # #         'new_files': [],
# # #         'updated_files': [],
# # #         'unchanged_files': [],
# # #         'deleted_files': [],
# # #         'new_folders': [],
# # #         'updated_folders': [],
# # #         'deleted_folders': [],
# # #         'skipped_files': []
# # #     }

# # #     def walk_folder(folder_id, parent_id=None, path=""):
# # #         query = f"'{folder_id}' in parents and trashed=false"
# # #         results = service.files().list(
# # #             q=query,
# # #             spaces='drive',
# # #             includeItemsFromAllDrives=True,
# # #             supportsAllDrives=True,
# # #             fields="files(id, name, mimeType, modifiedTime)"
# # #         ).execute()
# # #         items = results.get('files', [])

# # #         for item in items:
# # #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# # #             mime_type = None if node_type == "folder" else item['mimeType']
# # #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# # #             node_path = f"{path}.{safe_name}" if path else safe_name

# # #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# # #             current_nodes[item['id']] = {
# # #                 'name': item['name'],
# # #                 'parent_id': parent_id,
# # #                 'type': node_type,
# # #                 'mime_type': mime_type,
# # #                 'path': node_path,
# # #                 'modified_time': modified_time
# # #             }

# # #             if node_type == "folder":
# # #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# # #     walk_folder(root_folder_id)

# # #     # --- Postgres Sync (with folder tracking) ---
# # #     with conn.cursor() as cur:
# # #         # Get existing nodes (both files and folders) to track changes
# # #         cur.execute("SELECT id, type, modified_time FROM nodes")
# # #         existing_nodes = {row[0]: {'type': row[1], 'modified_time': row[2]} for row in cur.fetchall()}
        
# # #         for gid, info in current_nodes.items():
# # #             is_new = gid not in existing_nodes
# # #             is_updated = (not is_new and 
# # #                          existing_nodes[gid]['modified_time'] != info['modified_time'])
            
# # #             # Track changes by type
# # #             if is_new:
# # #                 if info['type'] == 'file':
# # #                     stats['new_files'].append(info['name'])
# # #                 else:
# # #                     stats['new_folders'].append(info['name'])
# # #             elif is_updated:
# # #                 if info['type'] == 'file':
# # #                     stats['updated_files'].append(info['name'])
# # #                 else:
# # #                     stats['updated_folders'].append(info['name'])
# # #             elif info['type'] == 'file':
# # #                 stats['unchanged_files'].append(info['name'])
            
# # #             cur.execute("""
# # #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# # #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# # #                 ON CONFLICT (id) DO UPDATE
# # #                 SET name = EXCLUDED.name,
# # #                     parent_id = EXCLUDED.parent_id,
# # #                     type = EXCLUDED.type,
# # #                     mime_type = EXCLUDED.mime_type,
# # #                     path = EXCLUDED.path,
# # #                     modified_time = EXCLUDED.modified_time
# # #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# # #             """, (gid, info['name'], info['parent_id'], info['type'],
# # #                   info['mime_type'], info['path'], info['modified_time']))

# # #         # Track deleted nodes
# # #         if current_nodes:
# # #             deleted_ids = set(existing_nodes.keys()) - set(current_nodes.keys())
# # #             if deleted_ids:
# # #                 cur.execute("SELECT name, type FROM nodes WHERE id IN %s", (tuple(deleted_ids),))
# # #                 for name, node_type in cur.fetchall():
# # #                     if node_type == 'file':
# # #                         stats['deleted_files'].append(name)
# # #                     else:
# # #                         stats['deleted_folders'].append(name)
            
# # #             cur.execute("DELETE FROM nodes WHERE id NOT IN %s", (tuple(current_nodes.keys()),))
# # #         else:
# # #             cur.execute("SELECT name, type FROM nodes")
# # #             for name, node_type in cur.fetchall():
# # #                 if node_type == 'file':
# # #                     stats['deleted_files'].append(name)
# # #                 else:
# # #                     stats['deleted_folders'].append(name)
# # #             cur.execute("DELETE FROM nodes;")
# # #         conn.commit()

# # #     # --- Qdrant Sync (INCREMENTAL - files only) ---
# # #     ensure_qdrant_collection()
    
# # #     try:
# # #         existing_points = qdrant.scroll(
# # #             collection_name=COLLECTION_NAME,
# # #             limit=10000,
# # #             with_payload=True,
# # #             with_vectors=False
# # #         )[0]
# # #         existing_drive_ids = {point.payload.get('drive_id'): {
# # #             'modified_time': point.payload.get('modified_time'),
# # #             'point_id': point.id
# # #         } for point in existing_points}
# # #     except Exception as e:
# # #         print(f"⚠️ Could not fetch existing Qdrant points: {e}")
# # #         existing_drive_ids = {}
    
# # #     files_to_add = []
# # #     files_to_update = []
# # #     files_to_delete = []
    
# # #     # Find files that need to be added or updated in Qdrant
# # #     for gid, info in current_nodes.items():
# # #         if info['type'] != "file" or not info['mime_type']:
# # #             continue
            
# # #         is_new = gid not in existing_drive_ids
# # #         is_updated = (not is_new and 
# # #                      existing_drive_ids[gid]['modified_time'] != info['modified_time'].isoformat())
        
# # #         if is_new:
# # #             files_to_add.append((gid, info))
# # #         elif is_updated:
# # #             files_to_update.append((gid, info))
    
# # #     # Find files in Qdrant that have been deleted from Drive
# # #     current_drive_file_ids = {gid for gid, info in current_nodes.items() if info['type'] == 'file'}
# # #     for drive_id, point_info in existing_drive_ids.items():
# # #         if drive_id not in current_drive_file_ids:
# # #             files_to_delete.append((drive_id, point_info['point_id']))
    
# # #     points_to_upsert = []
# # #     qdrant_stats = {
# # #         'new': [],
# # #         'updated': [],
# # #         'deleted': [],
# # #         'skipped': []
# # #     }
    
# # #     if files_to_add or files_to_update:
# # #         print(f"\n🔄 Processing {len(files_to_add)} new files and {len(files_to_update)} updated files for Qdrant...")
    
# # #     for gid, info in files_to_add + files_to_update:
# # #         is_new_file = (gid, info) in files_to_add
        
# # #         text_content = extract_file_content(service, gid, info['mime_type'], info['name'])
        
# # #         if text_content is None:
# # #             stats['skipped_files'].append(info['name'])
# # #             continue
        
# # #         if is_new_file:
# # #             qdrant_stats['new'].append(info['name'])
# # #         else:
# # #             qdrant_stats['updated'].append(info['name'])
        
# # #         point_id = drive_id_to_uuid(gid)
# # #         content_to_embed = text_content[:8000]
        
# # #         try:
# # #             embedding = generate_embedding(content_to_embed)
# # #         except Exception as e:
# # #             print(f"⚠️ Failed to generate embedding for {gid}: {e}")
# # #             continue
        
# # #         points_to_upsert.append(qmodels.PointStruct(
# # #             id=point_id,
# # #             vector=embedding,
# # #             payload={
# # #                 "drive_id": gid,
# # #                 "name": info['name'],
# # #                 "type": info['type'],
# # #                 "mime_type": info['mime_type'],
# # #                 "path": info['path'],
# # #                 "modified_time": info['modified_time'].isoformat(),
# # #                 "content": text_content[:10000]
# # #             }
# # #         ))
    
# # #     # Upload new/updated points to Qdrant in batches
# # #     if points_to_upsert:
# # #         total_points = len(points_to_upsert)
# # #         print(f"\n📤 Uploading {total_points} points to Qdrant in batches of {BATCH_SIZE}...")
        
# # #         for i in range(0, total_points, BATCH_SIZE):
# # #             batch = points_to_upsert[i:i + BATCH_SIZE]
# # #             try:
# # #                 qdrant.upsert(collection_name=COLLECTION_NAME, points=batch)
# # #                 print(f"✅ Uploaded batch {i // BATCH_SIZE + 1}/{(total_points + BATCH_SIZE - 1) // BATCH_SIZE} ({len(batch)} points)")
# # #             except Exception as e:
# # #                 print(f"❌ Failed to upload batch starting at index {i}: {e}")
# # #                 continue
        
# # #         print(f"✅ Successfully synced {total_points} items to Qdrant")
    
# # #     # Delete removed files from Qdrant
# # #     if files_to_delete:
# # #         print(f"\n🗑️  Deleting {len(files_to_delete)} removed files from Qdrant...")
# # #         point_ids_to_delete = [point_id for _, point_id in files_to_delete]
        
# # #         try:
# # #             qdrant.delete(
# # #                 collection_name=COLLECTION_NAME,
# # #                 points_selector=qmodels.PointIdsList(points=point_ids_to_delete)
# # #             )
# # #             qdrant_stats['deleted'] = [name for name in stats['deleted_files']]
# # #             print(f"✅ Deleted {len(point_ids_to_delete)} points from Qdrant")
# # #         except Exception as e:
# # #             print(f"❌ Failed to delete points from Qdrant: {e}")
    
# # #     # Print detailed statistics
# # #     print("\n" + "="*60)
# # #     print("📊 INCREMENTAL SYNC SUMMARY")
# # #     print("="*60)
    
# # #     print("\n🗄️  PostgreSQL Changes:")
    
# # #     # Files
# # #     if stats['new_files']:
# # #         print(f"  ✨ New files ({len(stats['new_files'])}):")
# # #         for name in stats['new_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['new_files']) > 10:
# # #             print(f"     ... and {len(stats['new_files']) - 10} more")
    
# # #     if stats['updated_files']:
# # #         print(f"  🔄 Updated files ({len(stats['updated_files'])}):")
# # #         for name in stats['updated_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['updated_files']) > 10:
# # #             print(f"     ... and {len(stats['updated_files']) - 10} more")
    
# # #     if stats['deleted_files']:
# # #         print(f"  🗑️  Deleted files ({len(stats['deleted_files'])}):")
# # #         for name in stats['deleted_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['deleted_files']) > 10:
# # #             print(f"     ... and {len(stats['deleted_files']) - 10} more")
    
# # #     if stats['unchanged_files']:
# # #         print(f"  ✓ Unchanged files: {len(stats['unchanged_files'])} (not processed)")
    
# # #     # Folders
# # #     if stats['new_folders']:
# # #         print(f"\n  📁 New folders ({len(stats['new_folders'])}):")
# # #         for name in stats['new_folders'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['new_folders']) > 10:
# # #             print(f"     ... and {len(stats['new_folders']) - 10} more")
    
# # #     if stats['updated_folders']:
# # #         print(f"  📁 Updated folders ({len(stats['updated_folders'])}):")
# # #         for name in stats['updated_folders'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['updated_folders']) > 10:
# # #             print(f"     ... and {len(stats['updated_folders']) - 10} more")
    
# # #     if stats['deleted_folders']:
# # #         print(f"  📁 Deleted folders ({len(stats['deleted_folders'])}):")
# # #         for name in stats['deleted_folders'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['deleted_folders']) > 10:
# # #             print(f"     ... and {len(stats['deleted_folders']) - 10} more")
    
# # #     print("\n🔍 Qdrant Vector Database Changes (Files Only):")
# # #     if qdrant_stats['new']:
# # #         print(f"  ✨ New embeddings ({len(qdrant_stats['new'])}):")
# # #         for name in qdrant_stats['new'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['new']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['new']) - 10} more")
    
# # #     if qdrant_stats['updated']:
# # #         print(f"  🔄 Updated embeddings ({len(qdrant_stats['updated'])}):")
# # #         for name in qdrant_stats['updated'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['updated']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['updated']) - 10} more")
    
# # #     if qdrant_stats['deleted']:
# # #         print(f"  🗑️  Deleted embeddings ({len(qdrant_stats['deleted'])}):")
# # #         for name in qdrant_stats['deleted'][:10]:
# # #             print(f"     • {name}")
# # #         if len(qdrant_stats['deleted']) > 10:
# # #             print(f"     ... and {len(qdrant_stats['deleted']) - 10} more")
    
# # #     if stats['skipped_files']:
# # #         print(f"\n⚠️  Skipped files ({len(stats['skipped_files'])}):")
# # #         for name in stats['skipped_files'][:10]:
# # #             print(f"     • {name}")
# # #         if len(stats['skipped_files']) > 10:
# # #             print(f"     ... and {len(stats['skipped_files']) - 10} more")
    
# # #     # Summary
# # #     total_pg_changes = (len(stats['new_files']) + len(stats['updated_files']) + 
# # #                        len(stats['deleted_files']) + len(stats['new_folders']) + 
# # #                        len(stats['updated_folders']) + len(stats['deleted_folders']))
# # #     total_qdrant_changes = (len(qdrant_stats['new']) + len(qdrant_stats['updated']) + 
# # #                            len(qdrant_stats['deleted']))
    
# # #     if total_pg_changes == 0 and total_qdrant_changes == 0:
# # #         print(f"\n  ✓ No changes detected - everything is up to date!")
    
# # #     print("\n" + "="*60)


# # # if __name__ == '__main__':
# # #     service = authenticate()
# # #     conn = psycopg2.connect(config.DATABASE_URL)

# # #     print("Syncing Google Drive → Postgres + Qdrant (INCREMENTAL MODE)...")
# # #     sync_drive(service, conn, config.ROOT_FOLDER_ID)
# # #     print("\nDone! Only changed files and folders were processed.")

# # #     conn.close()

# # from __future__ import print_function
# # import os
# # import io
# # import psycopg2
# # from urllib.parse import quote_plus
# # from google.oauth2.credentials import Credentials
# # from google_auth_oauthlib.flow import InstalledAppFlow
# # from google.auth.transport.requests import Request
# # from googleapiclient.discovery import build
# # from datetime import datetime
# # import uuid
# # import time

# # # Embedding
# # from openai import OpenAI

# # # PDF parsing
# # from PyPDF2 import PdfReader

# # # Document parsing
# # from docx import Document
# # import pytesseract
# # from PIL import Image

# # # Configuration
# # import config

# # # Google Drive scopes
# # SCOPES = [
# #     'https://www.googleapis.com/auth/drive.readonly',
# #     'https://www.googleapis.com/auth/drive.file'
# # ]

# # # Initialize OpenAI
# # openai_client = OpenAI(api_key=config.OPENAI_API_KEY)

# # # Embedding configuration
# # EMBEDDING_DIM = 1536
# # MAX_RETRIES = 3
# # RETRY_DELAY = 2  # seconds


# # # ----------------- Database Connection -----------------
# # def get_db_connection():
# #     """Create a new database connection with proper configuration."""
# #     try:
# #         conn = psycopg2.connect(
# #             config.DATABASE_URL,
# #             connect_timeout=30,
# #             keepalives=1,
# #             keepalives_idle=30,
# #             keepalives_interval=10,
# #             keepalives_count=5
# #         )
# #         # Set statement timeout to 5 minutes
# #         with conn.cursor() as cur:
# #             cur.execute("SET statement_timeout = '300000'")  # 5 minutes in ms
# #         conn.commit()
# #         return conn
# #     except Exception as e:
# #         print(f"Database connection failed: {e}")
# #         raise


# # def ensure_connection(conn):
# #     """Check if connection is alive, reconnect if needed."""
# #     try:
# #         with conn.cursor() as cur:
# #             cur.execute("SELECT 1")
# #         return conn
# #     except (psycopg2.OperationalError, psycopg2.InterfaceError):
# #         print("Connection lost, reconnecting...")
# #         return get_db_connection()


# # # ----------------- Authentication -----------------
# # def authenticate():
# #     creds = None
# #     if os.path.exists(config.TOKEN_FILE):
# #         creds = Credentials.from_authorized_user_file(config.TOKEN_FILE, SCOPES)
# #         print("Loaded existing token.json")

# #     if not creds or not creds.valid:
# #         if creds and creds.expired and creds.refresh_token:
# #             creds.refresh(Request())
# #             print("Token refreshed")
# #         else:
# #             flow = InstalledAppFlow.from_client_secrets_file(config.CREDENTIALS_FILE, SCOPES)
# #             creds = flow.run_local_server(port=0)

# #         with open(config.TOKEN_FILE, 'w') as token_file:
# #             token_file.write(creds.to_json())
# #         print("token.json saved successfully")

# #     return build('drive', 'v3', credentials=creds)


# # # ----------------- Embedding Helpers -----------------
# # def generate_embedding(text: str):
# #     """Generate embedding for text using OpenAI."""
# #     response = openai_client.embeddings.create(
# #         model="text-embedding-3-small",
# #         input=text
# #     )
# #     return response.data[0].embedding


# # def get_file_uuid(drive_id: str) -> str:
# #     """Generate a stable UUID for the file based on Drive ID."""
# #     return str(uuid.uuid5(uuid.NAMESPACE_DNS, f"drive.google.com/{drive_id}"))


# # # ----------------- File Extraction -----------------
# # def extract_file_content(service, file_id, mime_type, file_name=""):
# #     """Download and extract text content depending on file type."""
# #     try:
# #         # Google Docs
# #         if mime_type == "application/vnd.google-apps.document":
# #             doc = service.files().export(fileId=file_id, mimeType="text/plain").execute()
# #             return doc.decode("utf-8")

# #         # Google Sheets
# #         elif mime_type == "application/vnd.google-apps.spreadsheet":
# #             sheet = service.files().export(fileId=file_id, mimeType="text/csv").execute()
# #             return sheet.decode("utf-8")

# #         # PDF files
# #         elif mime_type == "application/pdf":
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             reader = PdfReader(fh)
# #             text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
# #             return text if text.strip() else None

# #         # Word Documents (.docx)
# #         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             doc = Document(fh)
# #             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
# #             return text if text.strip() else None

# #         # Images (JPEG, PNG, etc.) - use OCR
# #         elif mime_type.startswith("image/"):
# #             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
# #             image = Image.open(fh)
# #             text = pytesseract.image_to_string(image)
# #             return text.strip() if text.strip() else None

# #         # JSON files
# #         elif mime_type == "application/json":
# #             text = service.files().get_media(fileId=file_id).execute()
# #             content = text.decode("utf-8")
# #             return content if content.strip() else None

# #         # Plain text files
# #         elif mime_type.startswith("text/"):
# #             text = service.files().get_media(fileId=file_id).execute()
# #             content = text.decode("utf-8")
# #             return content if content.strip() else None

# #         else:
# #             print(f"Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
# #             return None

# #     except Exception as e:
# #         print(f"Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
# #         return None


# # # ----------------- PGVector Helpers -----------------
# # def upsert_embedding(conn, gid, info, text_content):
# #     """Insert or update a single embedding per file in pgvector table."""
# #     file_uuid = get_file_uuid(gid)
    
# #     # Use first 8000 characters for embedding (OpenAI's token limit consideration)
# #     content_to_embed = text_content[:8000]
    
# #     print(f"  Processing embedding for '{info['name']}'...")

# #     for attempt in range(MAX_RETRIES):
# #         try:
# #             conn = ensure_connection(conn)
            
# #             # Generate embedding
# #             try:
# #                 embedding = generate_embedding(content_to_embed)
# #             except Exception as e:
# #                 print(f"  Failed to generate embedding: {e}")
# #                 raise
            
# #             with conn.cursor() as cur:
# #                 # Single INSERT with ON CONFLICT on drive_id
# #                 cur.execute("""
# #                     INSERT INTO file_embeddings 
# #                     (id, drive_id, name, type, mime_type, path, modified_time, content, embedding)
# #                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
# #                     ON CONFLICT (drive_id) DO UPDATE
# #                     SET id = EXCLUDED.id,
# #                         name = EXCLUDED.name,
# #                         type = EXCLUDED.type,
# #                         mime_type = EXCLUDED.mime_type,
# #                         path = EXCLUDED.path,
# #                         modified_time = EXCLUDED.modified_time,
# #                         content = EXCLUDED.content,
# #                         embedding = EXCLUDED.embedding
# #                 """, (
# #                     file_uuid,
# #                     gid,
# #                     info['name'],
# #                     info['type'],
# #                     info['mime_type'],
# #                     info['path'],
# #                     info['modified_time'],
# #                     text_content[:10000],  # Store more content for reference
# #                     embedding
# #                 ))

# #                 # Update nodes table with vector reference
# #                 cur.execute("UPDATE nodes SET vector_id = %s WHERE id = %s", (file_uuid, gid))

# #             conn.commit()
# #             print(f"  Successfully embedded 1 chunk")
# #             return conn
            
# #         except (psycopg2.OperationalError, psycopg2.InterfaceError) as e:
# #             print(f"  Database error (attempt {attempt + 1}/{MAX_RETRIES}): {e}")
# #             if attempt < MAX_RETRIES - 1:
# #                 time.sleep(RETRY_DELAY * (attempt + 1))
# #                 conn = get_db_connection()
# #             else:
# #                 print(f"  Failed to embed '{info['name']}' after {MAX_RETRIES} attempts")
# #                 raise
# #         except Exception as e:
# #             print(f"  Unexpected error embedding '{info['name']}': {e}")
# #             conn.rollback()
# #             raise

# #     return conn


# # def delete_embedding(conn, drive_ids):
# #     """Remove embeddings for deleted files and reset vector_id in nodes."""
# #     try:
# #         conn = ensure_connection(conn)
# #         with conn.cursor() as cur:
# #             cur.execute("DELETE FROM file_embeddings WHERE drive_id = ANY(%s)", (drive_ids,))
# #             cur.execute("UPDATE nodes SET vector_id = NULL WHERE id = ANY(%s)", (drive_ids,))
# #         conn.commit()
# #         return conn
# #     except Exception as e:
# #         print(f"Error deleting embeddings: {e}")
# #         conn.rollback()
# #         raise


# # # ----------------- Google Drive Sync -----------------
# # def sync_drive(service, conn, root_folder_id):
# #     current_nodes = {}
# #     stats = {
# #         'new_files': [], 'updated_files': [], 'unchanged_files': [], 'deleted_files': [],
# #         'new_folders': [], 'updated_folders': [], 'deleted_folders': [], 'skipped_files': []
# #     }

# #     def walk_folder(folder_id, parent_id=None, path=""):
# #         query = f"'{folder_id}' in parents and trashed=false"
# #         results = service.files().list(
# #             q=query,
# #             spaces='drive',
# #             includeItemsFromAllDrives=True,
# #             supportsAllDrives=True,
# #             fields="files(id, name, mimeType, modifiedTime)"
# #         ).execute()
# #         items = results.get('files', [])

# #         for item in items:
# #             node_type = "folder" if item['mimeType'] == 'application/vnd.google-apps.folder' else "file"
# #             mime_type = None if node_type == "folder" else item['mimeType']
# #             safe_name = quote_plus(item['name'].replace(' ', '_'))
# #             node_path = f"{path}.{safe_name}" if path else safe_name

# #             modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

# #             current_nodes[item['id']] = {
# #                 'name': item['name'],
# #                 'parent_id': parent_id,
# #                 'type': node_type,
# #                 'mime_type': mime_type,
# #                 'path': node_path,
# #                 'modified_time': modified_time
# #             }

# #             if node_type == "folder":
# #                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

# #     walk_folder(root_folder_id)

# #     # --- Postgres Sync for nodes ---
# #     conn = ensure_connection(conn)
# #     with conn.cursor() as cur:
# #         cur.execute("SELECT id, type, modified_time FROM nodes")
# #         existing_nodes = {row[0]: {'type': row[1], 'modified_time': row[2]} for row in cur.fetchall()}

# #         for gid, info in current_nodes.items():
# #             is_new = gid not in existing_nodes
# #             is_updated = (not is_new and existing_nodes[gid]['modified_time'] != info['modified_time'])

# #             if is_new:
# #                 stats['new_files' if info['type'] == 'file' else 'new_folders'].append(info['name'])
# #             elif is_updated:
# #                 stats['updated_files' if info['type'] == 'file' else 'updated_folders'].append(info['name'])
# #             elif info['type'] == 'file':
# #                 stats['unchanged_files'].append(info['name'])

# #             cur.execute("""
# #                 INSERT INTO nodes (id, name, parent_id, type, mime_type, path, modified_time)
# #                 VALUES (%s, %s, %s, %s, %s, %s, %s)
# #                 ON CONFLICT (id) DO UPDATE
# #                 SET name = EXCLUDED.name,
# #                     parent_id = EXCLUDED.parent_id,
# #                     type = EXCLUDED.type,
# #                     mime_type = EXCLUDED.mime_type,
# #                     path = EXCLUDED.path,
# #                     modified_time = EXCLUDED.modified_time
# #                 WHERE nodes.modified_time IS DISTINCT FROM EXCLUDED.modified_time;
# #             """, (gid, info['name'], info['parent_id'], info['type'],
# #                   info['mime_type'], info['path'], info['modified_time']))

# #         # Detect deletions
# #         deleted_ids = set(existing_nodes.keys()) - set(current_nodes.keys())
# #         if deleted_ids:
# #             cur.execute("SELECT id, name, type FROM nodes WHERE id = ANY(%s)", (list(deleted_ids),))
# #             for node_id, name, node_type in cur.fetchall():
# #                 if node_type == 'file':
# #                     stats['deleted_files'].append(name)
# #                 else:
# #                     stats['deleted_folders'].append(name)
# #             cur.execute("DELETE FROM nodes WHERE id = ANY(%s)", (list(deleted_ids),))

# #         conn.commit()

# #     # --- Get existing embeddings to determine what needs processing ---
# #     with conn.cursor() as cur:
# #         cur.execute("SELECT drive_id, modified_time FROM file_embeddings")
# #         existing_embeddings = {row[0]: row[1] for row in cur.fetchall()}
    
# #     # --- Determine which files actually need embedding ---
# #     files_to_process = []
# #     embedding_stats = {
# #         'new_embeddings': [],
# #         'updated_embeddings': [],
# #         'unchanged_embeddings': []
# #     }
    
# #     for gid, info in current_nodes.items():
# #         if info['type'] != 'file':
# #             continue
        
# #         # Check if file needs embedding
# #         if gid not in existing_embeddings:
# #             # New file - needs embedding
# #             files_to_process.append((gid, info))
# #             embedding_stats['new_embeddings'].append(info['name'])
# #         elif existing_embeddings[gid] != info['modified_time']:
# #             # File was modified - needs re-embedding
# #             files_to_process.append((gid, info))
# #             embedding_stats['updated_embeddings'].append(info['name'])
# #         else:
# #             # File unchanged - skip embedding
# #             embedding_stats['unchanged_embeddings'].append(info['name'])
    
# #     if not files_to_process:
# #         print("\nNo files need embedding - all files are up to date!")
# #     else:
# #         print(f"\nProcessing {len(files_to_process)} files for embedding...")
# #         print(f"  - New: {len(embedding_stats['new_embeddings'])}")
# #         print(f"  - Updated: {len(embedding_stats['updated_embeddings'])}")
# #         print(f"  - Unchanged (skipped): {len(embedding_stats['unchanged_embeddings'])}")
    
# #     print(f"\nProcessing {len(files_to_process)} files for embedding...")
# #     for idx, (gid, info) in enumerate(files_to_process, 1):
# #         print(f"\n[{idx}/{len(files_to_process)}] Processing: {info['name']}")
# #         text_content = extract_file_content(service, gid, info['mime_type'], info['name'])
# #         if not text_content:
# #             stats['skipped_files'].append(info['name'])
# #             print(f"  Skipped (no content extracted)")
# #             continue
        
# #         try:
# #             conn = upsert_embedding(conn, gid, info, text_content)
# #         except Exception as e:
# #             print(f"  Failed to process '{info['name']}': {e}")
# #             stats['skipped_files'].append(info['name'])

# #     # --- Delete embeddings for removed files ---
# #     if deleted_ids:
# #         file_deleted_ids = [gid for gid in deleted_ids if gid in existing_nodes and existing_nodes[gid]['type'] == 'file']
# #         if file_deleted_ids:
# #             conn = delete_embedding(conn, file_deleted_ids)

# #     print("\n" + "="*60)
# #     print("SYNC SUMMARY")
# #     print("="*60)
    
# #     print("\nPostgreSQL Metadata Changes:")
# #     print(f"  New files: {len(stats['new_files'])}")
# #     print(f"  Updated files: {len(stats['updated_files'])}")
# #     print(f"  Deleted files: {len(stats['deleted_files'])}")
# #     print(f"  New folders: {len(stats['new_folders'])}")
# #     print(f"  Updated folders: {len(stats['updated_folders'])}")
# #     print(f"  Deleted folders: {len(stats['deleted_folders'])}")
    
# #     print("\nVector Embeddings (Actual Work Done):")
# #     print(f"  New embeddings created: {len(embedding_stats['new_embeddings'])}")
# #     print(f"  Embeddings updated: {len(embedding_stats['updated_embeddings'])}")
# #     print(f"  Unchanged (skipped): {len(embedding_stats['unchanged_embeddings'])}")
# #     print(f"  Skipped (no content): {len(stats['skipped_files'])}")
    
# #     if embedding_stats['new_embeddings']:
# #         print(f"\n  New files embedded:")
# #         for name in embedding_stats['new_embeddings'][:5]:
# #             print(f"    - {name}")
# #         if len(embedding_stats['new_embeddings']) > 5:
# #             print(f"    ... and {len(embedding_stats['new_embeddings']) - 5} more")
    
# #     if embedding_stats['updated_embeddings']:
# #         print(f"\n  Updated files re-embedded:")
# #         for name in embedding_stats['updated_embeddings'][:5]:
# #             print(f"    - {name}")
# #         if len(embedding_stats['updated_embeddings']) > 5:
# #             print(f"    ... and {len(embedding_stats['updated_embeddings']) - 5} more")
    
# #     print("="*60)

# #     return conn


# # # ----------------- Main -----------------
# # if __name__ == '__main__':
# #     conn = None
# #     try:
# #         service = authenticate()
# #         conn = get_db_connection()

# #         print("Syncing Google Drive → Supabase Postgres (pgvector)...")
# #         conn = sync_drive(service, conn, config.ROOT_FOLDER_ID)
# #         print("\nDone! Only changed files and folders were processed.")

# #     except KeyboardInterrupt:
# #         print("\nProcess interrupted by user")
# #     except Exception as e:
# #         print(f"\nFatal error: {e}")
# #         import traceback
# #         traceback.print_exc()
# #     finally:
# #         if conn:
# #             conn.close()
# #             print("Database connection closed")


# from __future__ import print_function
# import os
# import io
# import psycopg2
# from urllib.parse import quote_plus
# from google.oauth2.credentials import Credentials
# from google_auth_oauthlib.flow import InstalledAppFlow
# from google.auth.transport.requests import Request
# from googleapiclient.discovery import build
# from datetime import datetime
# import time
# import json
# from dotenv import load_dotenv

# # Embedding
# from openai import OpenAI

# # PDF parsing
# from PyPDF2 import PdfReader

# # Document parsing
# from docx import Document

# # Load environment variables
# load_dotenv()

# # Google Drive scopes
# SCOPES = [
#     'https://www.googleapis.com/auth/drive.readonly',
#     'https://www.googleapis.com/auth/drive.file'
# ]

# # Get configuration from environment
# OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
# DATABASE_URL = os.getenv('DATABASE_URL')
# ROOT_FOLDER_ID = os.getenv('ROOT_FOLDER_ID')
# CREDENTIALS_FILE = os.getenv('CREDENTIALS_FILE', 'credentials.json')
# TOKEN_FILE = os.getenv('TOKEN_FILE', 'token.json')

# # Initialize OpenAI
# openai_client = OpenAI(api_key=OPENAI_API_KEY)

# # Embedding configuration
# EMBEDDING_DIM = 1536
# MAX_RETRIES = 3
# RETRY_DELAY = 2  # seconds

# # Table name
# TABLE_NAME = "Haussler-Test-Data"

# # Supported file types - ONLY PDF and Word documents
# SUPPORTED_MIME_TYPES = {
#     'application/pdf',
#     'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
#     'application/msword',  # .doc
# }


# # ----------------- Database Connection -----------------
# def get_db_connection():
#     """Create a new database connection using Supabase IPv4-compatible pooler."""
#     try:
#         # Check if using pooler (IPv4 compatible) or direct connection
#         db_url = DATABASE_URL
        
#         if not db_url:
#             raise ValueError("DATABASE_URL not found in .env file")
        
#         # If using direct connection URL, warn user
#         if 'db.uohmhpinfcpsgiemyszk.supabase.co' in db_url:
#             print("WARNING: Using direct connection (IPv6 only)")
#             print("For IPv4 compatibility, update your .env DATABASE_URL to use the pooler:")
#             print("  Transaction Pooler: aws-1-ap-southeast-1.pooler.supabase.com:6543")
#             print("  Session Pooler: aws-1-ap-southeast-1.pooler.supabase.com:5432")
#             raise ConnectionError(
#                 "IPv6 connection detected. Please update DATABASE_URL in .env to use "
#                 "Supabase's IPv4-compatible pooler. See error message above for details."
#             )
        
#         print("Connecting to database via IPv4-compatible pooler...")
        
#         conn = psycopg2.connect(
#             db_url,
#             connect_timeout=30,
#             keepalives=1,
#             keepalives_idle=30,
#             keepalives_interval=10,
#             keepalives_count=5
#         )
        
#         # Set statement timeout to 5 minutes
#         with conn.cursor() as cur:
#             cur.execute("SET statement_timeout = '300000'")  # 5 minutes in ms
#         conn.commit()
        
#         print("Database connection established successfully")
#         return conn
        
#     except Exception as e:
#         print(f"Database connection failed: {e}")
#         raise


# def ensure_connection(conn):
#     """Check if connection is alive, reconnect if needed."""
#     try:
#         with conn.cursor() as cur:
#             cur.execute("SELECT 1")
#         return conn
#     except (psycopg2.OperationalError, psycopg2.InterfaceError):
#         print("Connection lost, reconnecting...")
#         return get_db_connection()


# # ----------------- Authentication -----------------
# def authenticate():
#     creds = None
#     if os.path.exists(TOKEN_FILE):
#         creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
#         print("Loaded existing token.json")

#     if not creds or not creds.valid:
#         if creds and creds.expired and creds.refresh_token:
#             creds.refresh(Request())
#             print("Token refreshed")
#         else:
#             flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
#             creds = flow.run_local_server(port=9003)

#         with open(TOKEN_FILE, 'w') as token_file:
#             token_file.write(creds.to_json())
#         print("token.json saved successfully")

#     return build('drive', 'v3', credentials=creds)


# # ----------------- Embedding Helpers -----------------
# def generate_embedding(text: str):
#     """Generate embedding for text using OpenAI."""
#     response = openai_client.embeddings.create(
#         model="text-embedding-3-small",
#         input=text
#     )
#     return response.data[0].embedding


# def get_drive_url(file_id: str, is_folder: bool = False) -> str:
#     """Generate Google Drive URL for a file or folder."""
#     if is_folder:
#         return f"https://drive.google.com/drive/folders/{file_id}"
#     else:
#         return f"https://drive.google.com/file/d/{file_id}/view"


# # ----------------- File Extraction -----------------
# def extract_file_content(service, file_id, mime_type, file_name=""):
#     """Download and extract text content - ONLY for PDF and Word documents."""
#     try:
#         # PDF files
#         if mime_type == "application/pdf":
#             print(f"  📄 Extracting PDF content...")
#             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
#             reader = PdfReader(fh)
            
#             # Extract text from all pages
#             text_parts = []
#             for page_num, page in enumerate(reader.pages, 1):
#                 try:
#                     page_text = page.extract_text()
#                     if page_text and page_text.strip():
#                         text_parts.append(page_text)
#                 except Exception as e:
#                     print(f"    ⚠ Warning: Could not extract page {page_num}: {e}")
#                     continue
            
#             text = "\n".join(text_parts)
            
#             if not text.strip():
#                 print(f"  ⚠ PDF appears to be empty or image-based (no extractable text)")
#                 return None
            
#             print(f"  ✓ Extracted {len(text)} characters from {len(text_parts)} pages")
#             return text

#         # Word Documents (.docx)
#         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             print(f"  📝 Extracting DOCX content...")
#             fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
#             doc = Document(fh)
            
#             # Extract all paragraphs
#             text_parts = []
#             for para in doc.paragraphs:
#                 if para.text.strip():
#                     text_parts.append(para.text)
            
#             # Also extract text from tables
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         if cell.text.strip():
#                             text_parts.append(cell.text)
            
#             text = "\n".join(text_parts)
            
#             if not text.strip():
#                 print(f"  ⚠ DOCX appears to be empty")
#                 return None
            
#             print(f"  ✓ Extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
#             return text

#         # Word Documents (.doc) - older format
#         elif mime_type == "application/msword":
#             print(f"  ⚠ .doc format detected - attempting basic extraction...")
#             try:
#                 # Try to read as binary and extract what we can
#                 fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
#                 # python-docx doesn't support .doc, so we'll skip these
#                 print(f"  ⚠ .doc format not fully supported. Please convert '{file_name}' to .docx for better results")
#                 return None
#             except Exception as e:
#                 print(f"  ✗ Failed to process .doc file: {e}")
#                 return None

#         else:
#             # This shouldn't happen if filtering is working correctly
#             print(f"  ⚠ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
#             return None

#     except Exception as e:
#         print(f"  ✗ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
#         import traceback
#         print(f"  Stack trace: {traceback.format_exc()}")
#         return None


# # ----------------- PGVector Helpers -----------------
# def upsert_embedding(conn, file_id, file_info, text_content):
#     """Insert or update a single embedding per file in pgvector table."""
    
#     # Use first 8000 characters for embedding (OpenAI's token limit consideration)
#     content_to_embed = text_content[:8000]
    
#     print(f"  Processing embedding for '{file_info['name']}'...")

#     for attempt in range(MAX_RETRIES):
#         try:
#             conn = ensure_connection(conn)
            
#             # Generate embedding
#             try:
#                 embedding = generate_embedding(content_to_embed)
#             except Exception as e:
#                 print(f"  Failed to generate embedding: {e}")
#                 raise
            
#             # Create metadata JSON
#             metadata = {
#                 "drive_id": file_id,
#                 "drive_url": get_drive_url(file_id),
#                 "name": file_info['name'],
#                 "mime_type": file_info['mime_type'],
#                 "modified_time": file_info['modified_time'].isoformat(),
#                 "parent_id": file_info.get('parent_id'),
#                 "path": file_info.get('path', '')
#             }
            
#             with conn.cursor() as cur:
#                 # Check if record exists based on drive_id in metadata
#                 cur.execute(f"""
#                     SELECT id FROM "{TABLE_NAME}" 
#                     WHERE metadata->>'drive_id' = %s
#                 """, (file_id,))
                
#                 existing_record = cur.fetchone()
                
#                 if existing_record:
#                     # Update existing record
#                     cur.execute(f"""
#                         UPDATE "{TABLE_NAME}"
#                         SET content = %s,
#                             embedding = %s,
#                             metadata = %s
#                         WHERE id = %s
#                     """, (
#                         text_content,
#                         embedding,
#                         json.dumps(metadata),
#                         existing_record[0]
#                     ))
#                 else:
#                     # Insert new record (id will be auto-generated)
#                     cur.execute(f"""
#                         INSERT INTO "{TABLE_NAME}" 
#                         (content, embedding, metadata)
#                         VALUES (%s, %s, %s)
#                     """, (
#                         text_content,
#                         embedding,
#                         json.dumps(metadata)
#                     ))

#             conn.commit()
#             print(f"  ✓ Successfully processed embedding")
#             return conn
            
#         except (psycopg2.OperationalError, psycopg2.InterfaceError) as e:
#             print(f"  Database error (attempt {attempt + 1}/{MAX_RETRIES}): {e}")
#             if attempt < MAX_RETRIES - 1:
#                 time.sleep(RETRY_DELAY * (attempt + 1))
#                 conn = get_db_connection()
#             else:
#                 print(f"  Failed to embed '{file_info['name']}' after {MAX_RETRIES} attempts")
#                 raise
#         except Exception as e:
#             print(f"  Unexpected error embedding '{file_info['name']}': {e}")
#             conn.rollback()
#             raise

#     return conn


# def delete_embeddings(conn, drive_ids):
#     """Remove embeddings for deleted files."""
#     try:
#         conn = ensure_connection(conn)
#         with conn.cursor() as cur:
#             for drive_id in drive_ids:
#                 cur.execute(f"""
#                     DELETE FROM "{TABLE_NAME}" 
#                     WHERE metadata->>'drive_id' = %s
#                 """, (drive_id,))
#         conn.commit()
#         return conn
#     except Exception as e:
#         print(f"Error deleting embeddings: {e}")
#         conn.rollback()
#         raise


# # ----------------- Google Drive Sync -----------------
# def sync_drive(service, conn, root_folder_id):
#     current_files = {}
#     stats = {
#         'new_files': [], 'updated_files': [], 'unchanged_files': [], 
#         'deleted_files': [], 'skipped_files': [], 'filtered_files': []
#     }

#     def walk_folder(folder_id, parent_id=None, path=""):
#         query = f"'{folder_id}' in parents and trashed=false"
#         results = service.files().list(
#             q=query,
#             spaces='drive',
#             includeItemsFromAllDrives=True,
#             supportsAllDrives=True,
#             fields="files(id, name, mimeType, modifiedTime)"
#         ).execute()
#         items = results.get('files', [])

#         for item in items:
#             is_folder = item['mimeType'] == 'application/vnd.google-apps.folder'
            
#             if not is_folder:
#                 # Filter: Only process supported file types
#                 if item['mimeType'] not in SUPPORTED_MIME_TYPES:
#                     stats['filtered_files'].append(item['name'])
#                     continue
                
#                 # Sanitize name for path
#                 safe_name = ''.join(c if c.isalnum() or c == '_' else '_' for c in item['name'])
#                 safe_name = '_'.join(filter(None, safe_name.split('_')))
#                 safe_name = safe_name.strip('_')
#                 if not safe_name:
#                     safe_name = 'unnamed'
                
#                 node_path = f"{path}.{safe_name}" if path else safe_name
#                 modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

#                 current_files[item['id']] = {
#                     'name': item['name'],
#                     'parent_id': parent_id,
#                     'mime_type': item['mimeType'],
#                     'path': node_path,
#                     'modified_time': modified_time
#                 }
#             else:
#                 # Recursively process folder contents
#                 safe_name = ''.join(c if c.isalnum() or c == '_' else '_' for c in item['name'])
#                 safe_name = '_'.join(filter(None, safe_name.split('_')))
#                 safe_name = safe_name.strip('_')
#                 if not safe_name:
#                     safe_name = 'unnamed'
                    
#                 node_path = f"{path}.{safe_name}" if path else safe_name
#                 walk_folder(item['id'], parent_id=item['id'], path=node_path)

#     print("Scanning Google Drive files (PDF, DOC, DOCX only)...")
#     walk_folder(root_folder_id)
#     print(f"Found {len(current_files)} supported files in Google Drive")
#     if stats['filtered_files']:
#         print(f"Filtered out {len(stats['filtered_files'])} unsupported file types")

#     # --- Get existing embeddings from database ---
#     conn = ensure_connection(conn)
#     with conn.cursor() as cur:
#         cur.execute(f"""
#             SELECT metadata->>'drive_id', metadata->>'modified_time' 
#             FROM "{TABLE_NAME}"
#         """)
#         existing_embeddings = {row[0]: row[1] for row in cur.fetchall()}
    
#     print(f"Found {len(existing_embeddings)} existing embeddings in database")
    
#     # --- Determine which files need processing ---
#     files_to_process = []
    
#     for file_id, file_info in current_files.items():
#         if file_id not in existing_embeddings:
#             # New file - needs embedding
#             files_to_process.append((file_id, file_info))
#             stats['new_files'].append(file_info['name'])
#         elif existing_embeddings[file_id] != file_info['modified_time'].isoformat():
#             # File was modified - needs re-embedding
#             files_to_process.append((file_id, file_info))
#             stats['updated_files'].append(file_info['name'])
#         else:
#             # File unchanged - skip embedding
#             stats['unchanged_files'].append(file_info['name'])
    
#     # --- Detect deletions ---
#     deleted_ids = set(existing_embeddings.keys()) - set(current_files.keys())
#     if deleted_ids:
#         print(f"\nFound {len(deleted_ids)} deleted files, removing from database...")
#         conn = delete_embeddings(conn, list(deleted_ids))
#         stats['deleted_files'] = list(deleted_ids)
    
#     # --- Process files that need embedding ---
#     if not files_to_process:
#         print("\n✓ No files need embedding - all files are up to date!")
#     else:
#         print(f"\n{'='*60}")
#         print(f"Processing {len(files_to_process)} files for embedding...")
#         print(f"  - New: {len(stats['new_files'])}")
#         print(f"  - Updated: {len(stats['updated_files'])}")
#         print(f"  - Unchanged (skipped): {len(stats['unchanged_files'])}")
#         print(f"{'='*60}\n")
        
#         for idx, (file_id, file_info) in enumerate(files_to_process, 1):
#             print(f"\n[{idx}/{len(files_to_process)}] Processing: {file_info['name']}")
#             print(f"  File type: {file_info['mime_type']}")
            
#             text_content = extract_file_content(service, file_id, file_info['mime_type'], file_info['name'])
            
#             if not text_content:
#                 stats['skipped_files'].append(file_info['name'])
#                 print(f"  ⚠ SKIPPED - No content extracted")
#                 print(f"  → This file may be empty, corrupted, or image-based PDF")
#                 continue
            
#             print(f"  → Content length: {len(text_content)} characters")
            
#             try:
#                 conn = upsert_embedding(conn, file_id, file_info, text_content)
#             except Exception as e:
#                 print(f"  ✗ Failed to process '{file_info['name']}': {e}")
#                 stats['skipped_files'].append(file_info['name'])

#     # --- Print Summary ---
#     print("\n" + "="*60)
#     print("SYNC SUMMARY")
#     print("="*60)
    
#     print(f"\n📊 Total supported files in Drive: {len(current_files)}")
#     print(f"📊 Files filtered (unsupported types): {len(stats['filtered_files'])}")
#     print(f"📊 Total embeddings in database: {len(existing_embeddings) - len(deleted_ids) + len(stats['new_files']) + len(stats['updated_files'])}")
    
#     print("\n📝 Changes:")
#     print(f"  ✓ New files embedded: {len(stats['new_files'])}")
#     print(f"  ↻ Files re-embedded: {len(stats['updated_files'])}")
#     print(f"  ✗ Files deleted: {len(stats['deleted_files'])}")
#     print(f"  ⊘ Files skipped (no content): {len(stats['skipped_files'])}")
#     print(f"  → Files unchanged: {len(stats['unchanged_files'])}")
    
#     if stats['new_files']:
#         print(f"\n  ✓ New files embedded:")
#         for name in stats['new_files'][:5]:
#             print(f"    • {name}")
#         if len(stats['new_files']) > 5:
#             print(f"    ... and {len(stats['new_files']) - 5} more")
    
#     if stats['updated_files']:
#         print(f"\n  ↻ Updated files re-embedded:")
#         for name in stats['updated_files'][:5]:
#             print(f"    • {name}")
#         if len(stats['updated_files']) > 5:
#             print(f"    ... and {len(stats['updated_files']) - 5} more")
    
#     if stats['skipped_files']:
#         print(f"\n  ⚠ SKIPPED FILES (no content extracted):")
#         for name in stats['skipped_files'][:10]:
#             print(f"    • {name}")
#         if len(stats['skipped_files']) > 10:
#             print(f"    ... and {len(stats['skipped_files']) - 10} more")
#         print(f"\n  💡 Common reasons for skipped files:")
#         print(f"    - Image-based PDFs (scanned documents without OCR)")
#         print(f"    - Empty or corrupted files")
#         print(f"    - Old .doc format (convert to .docx)")
#         print(f"    - Password-protected documents")
    
#     print("="*60)

#     return conn


# # ----------------- Main -----------------
# if __name__ == '__main__':
#     conn = None
#     try:
#         print("="*60)
#         print("Google Drive → Haussler Database Sync")
#         print("Processing: PDF, DOC, DOCX files only")
#         print("="*60)
        
#         # Validate environment variables
#         if not OPENAI_API_KEY:
#             raise ValueError("OPENAI_API_KEY not found in .env file")
#         if not ROOT_FOLDER_ID:
#             raise ValueError("ROOT_FOLDER_ID not found in .env file")
        
#         service = authenticate()
#         conn = get_db_connection()

#         print(f"\nSyncing to table: {TABLE_NAME}")
#         conn = sync_drive(service, conn, ROOT_FOLDER_ID)
#         print("\n✓ Done! Only changed files were processed.")

#     except KeyboardInterrupt:
#         print("\n\n⚠ Process interrupted by user")
#     except Exception as e:
#         print(f"\n✗ Fatal error: {e}")
#         import traceback
#         traceback.print_exc()
#     finally:
#         if conn:
#             conn.close()
#             print("Database connection closed")

from __future__ import print_function
import os
import io
import psycopg2
from urllib.parse import quote_plus
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from datetime import datetime
import time
import json
from dotenv import load_dotenv

# Embedding
from openai import OpenAI

# PDF parsing
from PyPDF2 import PdfReader

# Document parsing
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Google Drive scopes - READ ONLY (never modifies Drive)
SCOPES = [
    'https://www.googleapis.com/auth/drive.readonly',  # READ ONLY - Safe!
]

# Get configuration from environment
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
DATABASE_URL = os.getenv('DATABASE_URL')
ROOT_FOLDER_ID = os.getenv('ROOT_FOLDER_ID')
CREDENTIALS_FILE = os.getenv('CREDENTIALS_FILE', 'credentials.json')
TOKEN_FILE = os.getenv('TOKEN_FILE', 'token.json')

# Validate required environment variables
def validate_env_vars():
    """Check that all required environment variables are set."""
    missing_vars = []
    
    if not OPENAI_API_KEY:
        missing_vars.append('OPENAI_API_KEY')
    if not DATABASE_URL:
        missing_vars.append('DATABASE_URL')
    if not ROOT_FOLDER_ID:
        missing_vars.append('ROOT_FOLDER_ID')
    
    if missing_vars:
        print("\n❌ ERROR: Missing required environment variables in .env file:")
        for var in missing_vars:
            print(f"   - {var}")
        print("\n📝 Please create a .env file in the same directory as this script with:")
        print("   OPENAI_API_KEY=your_key_here")
        print("   DATABASE_URL=your_db_url_here")
        print("   ROOT_FOLDER_ID=your_folder_id_here")
        print("\n💡 Make sure the .env file is in: " + os.getcwd())
        raise ValueError(f"Missing environment variables: {', '.join(missing_vars)}")
    
    print("✓ All environment variables loaded successfully")

# Validate before initializing OpenAI
validate_env_vars()

# Initialize OpenAI
openai_client = OpenAI(api_key=OPENAI_API_KEY)

# Embedding configuration
EMBEDDING_DIM = 1536
MAX_RETRIES = 3
RETRY_DELAY = 2  # seconds

# Table name
TABLE_NAME = "Haussler-Test-Data"

# Supported file types - ONLY PDF and Word documents
SUPPORTED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
    'application/msword',  # .doc
}


# ----------------- Database Connection -----------------
def get_db_connection():
    """Create a new database connection using Supabase IPv4-compatible pooler."""
    try:
        # Check if using pooler (IPv4 compatible) or direct connection
        db_url = DATABASE_URL
        
        if not db_url:
            raise ValueError("DATABASE_URL not found in .env file")
        
        # If using direct connection URL, warn user
        if 'db.uohmhpinfcpsgiemyszk.supabase.co' in db_url:
            print("WARNING: Using direct connection (IPv6 only)")
            print("For IPv4 compatibility, update your .env DATABASE_URL to use the pooler:")
            print("  Transaction Pooler: aws-1-ap-southeast-1.pooler.supabase.com:6543")
            print("  Session Pooler: aws-1-ap-southeast-1.pooler.supabase.com:5432")
            raise ConnectionError(
                "IPv6 connection detected. Please update DATABASE_URL in .env to use "
                "Supabase's IPv4-compatible pooler. See error message above for details."
            )
        
        print("Connecting to database via IPv4-compatible pooler...")
        
        conn = psycopg2.connect(
            db_url,
            connect_timeout=30,
            keepalives=1,
            keepalives_idle=30,
            keepalives_interval=10,
            keepalives_count=5
        )
        
        # Set statement timeout to 5 minutes
        with conn.cursor() as cur:
            cur.execute("SET statement_timeout = '300000'")  # 5 minutes in ms
        conn.commit()
        
        print("Database connection established successfully")
        return conn
        
    except Exception as e:
        print(f"Database connection failed: {e}")
        raise


def ensure_connection(conn):
    """Check if connection is alive, reconnect if needed."""
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT 1")
        return conn
    except (psycopg2.OperationalError, psycopg2.InterfaceError):
        print("Connection lost, reconnecting...")
        return get_db_connection()


# ----------------- Authentication -----------------
def authenticate():
    creds = None
    if os.path.exists(TOKEN_FILE):
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
        print("Loaded existing token.json")

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
            print("Token refreshed")
        else:
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
            creds = flow.run_local_server(port=9003)

        with open(TOKEN_FILE, 'w') as token_file:
            token_file.write(creds.to_json())
        print("token.json saved successfully")

    return build('drive', 'v3', credentials=creds)


# ----------------- Embedding Helpers -----------------
def generate_embedding(text: str):
    """Generate embedding for text using OpenAI."""
    response = openai_client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return response.data[0].embedding


def get_drive_url(file_id: str, is_folder: bool = False) -> str:
    """Generate Google Drive URL for a file or folder."""
    if is_folder:
        return f"https://drive.google.com/drive/folders/{file_id}"
    else:
        return f"https://drive.google.com/file/d/{file_id}/view"


# ----------------- File Extraction -----------------
def extract_file_content(service, file_id, mime_type, file_name=""):
    """Download and extract text content - ONLY for PDF and Word documents."""
    try:
        # PDF files
        if mime_type == "application/pdf":
            print(f"  📄 Extracting PDF content...")
            fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
            reader = PdfReader(fh)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    print(f"    ⚠ Warning: Could not extract page {page_num}: {e}")
                    continue
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                print(f"  ⚠ PDF appears to be empty or image-based (no extractable text)")
                return None
            
            print(f"  ✓ Extracted {len(text)} characters from {len(text_parts)} pages")
            return text

        # Word Documents (.docx)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            print(f"  📝 Extracting DOCX content...")
            fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
            doc = Document(fh)
            
            # Extract all paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                print(f"  ⚠ DOCX appears to be empty")
                return None
            
            print(f"  ✓ Extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
            return text

        # Word Documents (.doc) - older format
        elif mime_type == "application/msword":
            print(f"  ⚠ .doc format detected - attempting basic extraction...")
            try:
                # Try to read as binary and extract what we can
                fh = io.BytesIO(service.files().get_media(fileId=file_id).execute())
                # python-docx doesn't support .doc, so we'll skip these
                print(f"  ⚠ .doc format not fully supported. Please convert '{file_name}' to .docx for better results")
                return None
            except Exception as e:
                print(f"  ✗ Failed to process .doc file: {e}")
                return None

        else:
            # This shouldn't happen if filtering is working correctly
            print(f"  ⚠ Unsupported MIME type '{mime_type}' for file '{file_name}' ({file_id})")
            return None

    except Exception as e:
        print(f"  ✗ Could not extract content from '{file_name}' ({file_id}): {type(e).__name__}: {e}")
        import traceback
        print(f"  Stack trace: {traceback.format_exc()}")
        return None


# ----------------- PGVector Helpers -----------------
def upsert_embedding(conn, file_id, file_info, text_content):
    """Insert or update a single embedding per file in pgvector table."""
    
    # Use first 8000 characters for embedding (OpenAI's token limit consideration)
    content_to_embed = text_content[:8000]
    
    print(f"  Processing embedding for '{file_info['name']}'...")

    for attempt in range(MAX_RETRIES):
        try:
            conn = ensure_connection(conn)
            
            # Generate embedding
            try:
                embedding = generate_embedding(content_to_embed)
            except Exception as e:
                print(f"  Failed to generate embedding: {e}")
                raise
            
            # Create metadata JSON
            metadata = {
                "drive_id": file_id,
                "drive_url": get_drive_url(file_id),
                "name": file_info['name'],
                "mime_type": file_info['mime_type'],
                "modified_time": file_info['modified_time'].isoformat(),
                "parent_id": file_info.get('parent_id'),
                "path": file_info.get('path', '')
            }
            
            with conn.cursor() as cur:
                # Check if record exists based on drive_id in metadata
                cur.execute(f"""
                    SELECT id FROM "{TABLE_NAME}" 
                    WHERE metadata->>'drive_id' = %s
                """, (file_id,))
                
                existing_record = cur.fetchone()
                
                if existing_record:
                    # Update existing record
                    cur.execute(f"""
                        UPDATE "{TABLE_NAME}"
                        SET content = %s,
                            embedding = %s,
                            metadata = %s
                        WHERE id = %s
                    """, (
                        text_content,
                        embedding,
                        json.dumps(metadata),
                        existing_record[0]
                    ))
                else:
                    # Insert new record (id will be auto-generated)
                    cur.execute(f"""
                        INSERT INTO "{TABLE_NAME}" 
                        (content, embedding, metadata)
                        VALUES (%s, %s, %s)
                    """, (
                        text_content,
                        embedding,
                        json.dumps(metadata)
                    ))

            conn.commit()
            print(f"  ✓ Successfully processed embedding")
            return conn
            
        except (psycopg2.OperationalError, psycopg2.InterfaceError) as e:
            print(f"  Database error (attempt {attempt + 1}/{MAX_RETRIES}): {e}")
            if attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY * (attempt + 1))
                conn = get_db_connection()
            else:
                print(f"  Failed to embed '{file_info['name']}' after {MAX_RETRIES} attempts")
                raise
        except Exception as e:
            print(f"  Unexpected error embedding '{file_info['name']}': {e}")
            conn.rollback()
            raise

    return conn


def delete_embeddings(conn, drive_ids):
    """Remove embeddings for deleted files."""
    try:
        conn = ensure_connection(conn)
        with conn.cursor() as cur:
            for drive_id in drive_ids:
                cur.execute(f"""
                    DELETE FROM "{TABLE_NAME}" 
                    WHERE metadata->>'drive_id' = %s
                """, (drive_id,))
        conn.commit()
        return conn
    except Exception as e:
        print(f"Error deleting embeddings: {e}")
        conn.rollback()
        raise


# ----------------- Google Drive Sync -----------------
def sync_drive(service, conn, root_folder_id):
    current_files = {}
    stats = {
        'new_files': [], 'updated_files': [], 'unchanged_files': [], 
        'deleted_files': [], 'skipped_files': [], 'filtered_files': []
    }

    def walk_folder(folder_id, parent_id=None, path=""):
        query = f"'{folder_id}' in parents and trashed=false"
        results = service.files().list(
            q=query,
            spaces='drive',
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
            fields="files(id, name, mimeType, modifiedTime)"
        ).execute()
        items = results.get('files', [])

        for item in items:
            is_folder = item['mimeType'] == 'application/vnd.google-apps.folder'
            
            if not is_folder:
                # Filter: Only process supported file types
                if item['mimeType'] not in SUPPORTED_MIME_TYPES:
                    stats['filtered_files'].append(item['name'])
                    continue
                
                # Sanitize name for path
                safe_name = ''.join(c if c.isalnum() or c == '_' else '_' for c in item['name'])
                safe_name = '_'.join(filter(None, safe_name.split('_')))
                safe_name = safe_name.strip('_')
                if not safe_name:
                    safe_name = 'unnamed'
                
                node_path = f"{path}.{safe_name}" if path else safe_name
                modified_time = datetime.fromisoformat(item['modifiedTime'].replace("Z", "+00:00"))

                current_files[item['id']] = {
                    'name': item['name'],
                    'parent_id': parent_id,
                    'mime_type': item['mimeType'],
                    'path': node_path,
                    'modified_time': modified_time
                }
            else:
                # Recursively process folder contents
                safe_name = ''.join(c if c.isalnum() or c == '_' else '_' for c in item['name'])
                safe_name = '_'.join(filter(None, safe_name.split('_')))
                safe_name = safe_name.strip('_')
                if not safe_name:
                    safe_name = 'unnamed'
                    
                node_path = f"{path}.{safe_name}" if path else safe_name
                walk_folder(item['id'], parent_id=item['id'], path=node_path)

    print("Scanning Google Drive files (PDF, DOC, DOCX only)...")
    walk_folder(root_folder_id)
    print(f"Found {len(current_files)} supported files in Google Drive")
    if stats['filtered_files']:
        print(f"Filtered out {len(stats['filtered_files'])} unsupported file types")

    # --- Get existing embeddings from database ---
    conn = ensure_connection(conn)
    with conn.cursor() as cur:
        cur.execute(f"""
            SELECT metadata->>'drive_id', metadata->>'modified_time' 
            FROM "{TABLE_NAME}"
        """)
        existing_embeddings = {row[0]: row[1] for row in cur.fetchall()}
    
    print(f"Found {len(existing_embeddings)} existing embeddings in database")
    
    # --- Determine which files need processing ---
    files_to_process = []
    
    for file_id, file_info in current_files.items():
        if file_id not in existing_embeddings:
            # New file - needs embedding
            files_to_process.append((file_id, file_info))
            stats['new_files'].append(file_info['name'])
        elif existing_embeddings[file_id] != file_info['modified_time'].isoformat():
            # File was modified - needs re-embedding
            files_to_process.append((file_id, file_info))
            stats['updated_files'].append(file_info['name'])
        else:
            # File unchanged - skip embedding
            stats['unchanged_files'].append(file_info['name'])
    
    # --- Detect deletions ---
    deleted_ids = set(existing_embeddings.keys()) - set(current_files.keys())
    if deleted_ids:
        print(f"\nFound {len(deleted_ids)} deleted files, removing from database...")
        conn = delete_embeddings(conn, list(deleted_ids))
        stats['deleted_files'] = list(deleted_ids)
    
    # --- Process files that need embedding ---
    if not files_to_process:
        print("\n✓ No files need embedding - all files are up to date!")
    else:
        print(f"\n{'='*60}")
        print(f"Processing {len(files_to_process)} files for embedding...")
        print(f"  - New: {len(stats['new_files'])}")
        print(f"  - Updated: {len(stats['updated_files'])}")
        print(f"  - Unchanged (skipped): {len(stats['unchanged_files'])}")
        print(f"{'='*60}\n")
        
        for idx, (file_id, file_info) in enumerate(files_to_process, 1):
            print(f"\n[{idx}/{len(files_to_process)}] Processing: {file_info['name']}")
            print(f"  File type: {file_info['mime_type']}")
            
            text_content = extract_file_content(service, file_id, file_info['mime_type'], file_info['name'])
            
            if not text_content:
                stats['skipped_files'].append(file_info['name'])
                print(f"  ⚠ SKIPPED - No content extracted")
                print(f"  → This file may be empty, corrupted, or image-based PDF")
                continue
            
            print(f"  → Content length: {len(text_content)} characters")
            
            try:
                conn = upsert_embedding(conn, file_id, file_info, text_content)
            except Exception as e:
                print(f"  ✗ Failed to process '{file_info['name']}': {e}")
                stats['skipped_files'].append(file_info['name'])

    # --- Print Summary ---
    print("\n" + "="*60)
    print("SYNC SUMMARY")
    print("="*60)
    
    print(f"\n📊 Total supported files in Drive: {len(current_files)}")
    print(f"📊 Files filtered (unsupported types): {len(stats['filtered_files'])}")
    print(f"📊 Total embeddings in database: {len(existing_embeddings) - len(deleted_ids) + len(stats['new_files']) + len(stats['updated_files'])}")
    
    print("\n📝 Changes:")
    print(f"  ✓ New files embedded: {len(stats['new_files'])}")
    print(f"  ↻ Files re-embedded: {len(stats['updated_files'])}")
    print(f"  ✗ Files deleted: {len(stats['deleted_files'])}")
    print(f"  ⊘ Files skipped (no content): {len(stats['skipped_files'])}")
    print(f"  → Files unchanged: {len(stats['unchanged_files'])}")
    
    if stats['new_files']:
        print(f"\n  ✓ New files embedded:")
        for name in stats['new_files'][:5]:
            print(f"    • {name}")
        if len(stats['new_files']) > 5:
            print(f"    ... and {len(stats['new_files']) - 5} more")
    
    if stats['updated_files']:
        print(f"\n  ↻ Updated files re-embedded:")
        for name in stats['updated_files'][:5]:
            print(f"    • {name}")
        if len(stats['updated_files']) > 5:
            print(f"    ... and {len(stats['updated_files']) - 5} more")
    
    if stats['skipped_files']:
        print(f"\n  ⚠ SKIPPED FILES (no content extracted):")
        for name in stats['skipped_files'][:10]:
            print(f"    • {name}")
        if len(stats['skipped_files']) > 10:
            print(f"    ... and {len(stats['skipped_files']) - 10} more")
        print(f"\n  💡 Common reasons for skipped files:")
        print(f"    - Image-based PDFs (scanned documents without OCR)")
        print(f"    - Empty or corrupted files")
        print(f"    - Old .doc format (convert to .docx)")
        print(f"    - Password-protected documents")
    
    print("="*60)

    return conn


# ----------------- Main -----------------
if __name__ == '__main__':
    conn = None
    try:
        print("="*60)
        print("Google Drive → Haussler Database Sync")
        print("Processing: PDF, DOC, DOCX files only")
        print("🔒 READ-ONLY MODE: Drive files will NEVER be modified")
        print("="*60)
        
        print(f"\n📁 Current directory: {os.getcwd()}")
        print(f"🔍 Looking for .env file...")
        
        if not os.path.exists('.env'):
            print("\n❌ ERROR: .env file not found!")
            print("📝 Please create a .env file in the current directory with your credentials")
            print("\nExample .env file:")
            print("-" * 40)
            print("OPENAI_API_KEY=sk-...")
            print("DATABASE_URL=postgresql://...")
            print("ROOT_FOLDER_ID=1gRGcru...")
            print("CREDENTIALS_FILE=credentials.json")
            print("TOKEN_FILE=token.json")
            print("-" * 40)
            exit(1)
        
        print("✓ .env file found")
        
        # This will validate and show error if variables are missing
        validate_env_vars()
        
        service = authenticate()
        conn = get_db_connection()

        print(f"\nSyncing to table: {TABLE_NAME}")
        conn = sync_drive(service, conn, ROOT_FOLDER_ID)
        print("\n✓ Done! Only changed files were processed.")

    except KeyboardInterrupt:
        print("\n\n⚠ Process interrupted by user")
    except Exception as e:
        print(f"\n✗ Fatal error: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if conn:
            conn.close()
            print("Database connection closed")
